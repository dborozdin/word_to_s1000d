#!/usr/bin/env python3
"""
Детальный тест для проверки многолистовых иллюстраций.
"""

import sys
import os
import tempfile
from docx import Document

# Добавляем путь к модулям проекта
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from parsers.multi_sheet_illustration_parser import process_multi_sheet_illustrations, find_multi_sheet_references
from processing_scripts.descriptive_processor import process_descriptive_document
import re

def detailed_test():
    """Детальный тест с проверкой XML."""
    print("🚀 Детальное тестирование многолистовых иллюстраций")
    
    # Создаем документ с различными типами иллюстраций
    doc = Document()
    doc.add_heading('Тест многолистовых иллюстраций', 0)
    doc.add_paragraph('Этот документ содержит различные типы иллюстраций для тестирования.')
    
    # Обычная иллюстрация
    doc.add_heading('Обычная иллюстрация', level=1)
    doc.add_paragraph('Рисунок 1 - Обычная иллюстрация без листов')
    
    # Многолистовая иллюстрация
    doc.add_heading('Многолистовая иллюстрация', level=1)
    doc.add_paragraph('Рисунок 2 - Схема системы (лист 1, лист 2)')
    
    # Еще одна многолистовая
    doc.add_heading('Еще одна многолистовая', level=1)
    doc.add_paragraph('Рисунок 3 - Алгоритм работы (лист 1, лист 2, лист 3)')
    
    # Сохраняем во временный файл
    temp_dir = tempfile.mkdtemp()
    test_file = os.path.join(temp_dir, 'detailed_test.docx')
    doc.save(test_file)
    
    print(f"Создан тестовый файл: {test_file}")
    
    # Обрабатываем документ
    output_dir = tempfile.mkdtemp()
    try:
        print(f"Выходная директория: {output_dir}")
        files = process_descriptive_document(test_file, output_dir)
        
        print(f"\nСоздано файлов: {len(files)}")
        for filepath in files:
            print(f"  - {os.path.basename(filepath)}")
        
        # Анализируем XML
        for xml_file in files:
            if xml_file.endswith('.xml') and 'DMC' in os.path.basename(xml_file):
                print(f"\n=== АНАЛИЗ XML: {os.path.basename(xml_file)} ===")
                
                with open(xml_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # Ищем все figure элементы
                figure_pattern = r'<figure[^>]*id="([^"]*)"[^>]*>.*?<title>(.*?)</title>(.*?)</figure>'
                figures = re.findall(figure_pattern, content, re.DOTALL | re.IGNORECASE)
                
                print(f"Найдено figure элементов: {len(figures)}")
                
                multi_sheet_count = 0
                regular_count = 0
                
                for fig_id, title, fig_content in figures:
                    # Подсчитываем graphic элементы
                    graphic_count = len(re.findall(r'<graphic[^>]*>', fig_content, re.IGNORECASE))
                    
                    if graphic_count > 1:
                        multi_sheet_count += 1
                        print(f"\n📊 МНОГОЛИСТОВАЯ ИЛЛЮСТРАЦИЯ:")
                        print(f"   ID: {fig_id}")
                        print(f"   Заголовок: {title.strip()}")
                        print(f"   Количество листов: {graphic_count}")
                        
                        # Выводим детали каждого листа
                        graphics = re.findall(r'<graphic[^>]*infoEntityIdent="([^"]*)"[^>]*>', fig_content, re.IGNORECASE)
                        for i, graphic_ident in enumerate(graphics, 1):
                            print(f"     Лист {i}: {graphic_ident}")
                    
                    else:
                        regular_count += 1
                        print(f"\n📷 ОБЫЧНАЯ ИЛЛЮСТРАЦИЯ:")
                        print(f"   ID: {fig_id}")
                        print(f"   Заголовок: {title.strip()}")
                        print(f"   Количество листов: {graphic_count}")
                
                print(f"\n=== ИТОГИ ===")
                print(f"Многолистовых иллюстраций: {multi_sheet_count}")
                print(f"Обычных иллюстраций: {regular_count}")
                
                # Проверяем результаты
                expected_multi_sheet = 2  # Рисунки 2 и 3
                expected_regular = 1      # Рисунок 1
                
                if multi_sheet_count == expected_multi_sheet and regular_count == expected_regular:
                    print("✅ ТЕСТ ПРОЙДЕН: Все иллюстрации обработаны правильно!")
                    return True
                else:
                    print(f"❌ ТЕСТ НЕ ПРОЙДЕН: ожидалось {expected_multi_sheet} многолистовых и {expected_regular} обычных")
                    return False
                
    except Exception as e:
        print(f"❌ Ошибка: {e}")
        import traceback
        traceback.print_exc()
        return False
        
    finally:
        import shutil
        shutil.rmtree(temp_dir)
        shutil.rmtree(output_dir)

if __name__ == "__main__":
    success = detailed_test()
    if success:
        print("\n🎉 ВСЕ ТЕСТЫ ПРОЙДЕНЫ УСПЕШНО!")
    else:
        print("\n💥 ТЕСТЫ НЕ ПРОЙДЕНЫ")
