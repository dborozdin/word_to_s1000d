#!/usr/bin/env python3
"""
Script to analyze numbered paragraph headers in the test document РСУ_адаптированная.docx
"""

import os
import re
from docx import Document

def analyze_numbered_headers(doc_path: str, output_log: str):
    """
    Analyze numbered paragraph headers in document and write results to log file.

    Args:
        doc_path: Path to the document
        output_log: Path for output log file
    """
    print(f"Analyzing numbered headers in document: {doc_path}")

    # Check if document exists
    if not os.path.exists(doc_path):
        error_msg = f"Document not found: {doc_path}"
        print(error_msg)
        with open(output_log, 'w', encoding='utf-8') as f:
            f.write(f"ERROR: {error_msg}\n")
        return

    try:
        # Load document
        doc = Document(doc_path)
        print(f"Document loaded successfully. Total paragraphs: {len(doc.paragraphs)}")

        # Find numbered headers - first let's see ALL paragraphs that start with numbers
        numbered_paragraphs = []

        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()

            if not text:
                continue

            # Check if this starts with a number pattern
            if re.match(r'^\s*\d+\.\s*\S+', text):
                numbered_paragraphs.append({
                    'para_idx': para_idx,
                    'text': text,
                    'style': paragraph.style.name if paragraph.style else 'Normal',
                    'has_continuation': para_idx + 1 < len(doc.paragraphs) and
                                      doc.paragraphs[para_idx + 1].text.strip() and
                                      not re.match(r'^\s*\d+\.\s*', doc.paragraphs[para_idx + 1].text.strip())
                })

        print(f"Found {len(numbered_paragraphs)} paragraphs starting with numbers")

        # Now analyze which ones look like headers vs list items
        numbered_headers = []
        list_items = []

        for para in numbered_paragraphs:
            text = para['text']
            lines = text.split('\n')

            # Heuristic: if it has multiple lines, first line is short, and followed by content
            # OR if it's a single short line that looks like a title
            is_header = False

            if len(lines) > 1:
                # Multi-line: check if first line is short and looks like title
                first_line = lines[0].strip()
                if len(first_line) < 100 and not first_line.endswith((':', ';', '.')):
                    is_header = True
            elif len(text) < 100 and not text.endswith((':', ';')):
                # Single line: short and doesn't end with list punctuation
                is_header = True

            if is_header:
                numbered_headers.append(para)
            else:
                list_items.append(para)

        print(f"Found {len(numbered_headers)} potential numbered headers")

        # Write log
        with open(output_log, 'w', encoding='utf-8') as f:
            f.write("Анализ нумерованных заголовков параграфов\n")
            f.write("=" * 60 + "\n")
            f.write(f"Документ: {os.path.basename(doc_path)}\n")
            f.write(f"Всего параграфов: {len(doc.paragraphs)}\n")
            f.write(f"Найдено нумерованных заголовков: {len(numbered_headers)}\n\n")

            if not numbered_headers:
                f.write("Нумерованные заголовки параграфов не найдены.\n")
                return

            for i, header in enumerate(numbered_headers, 1):
                f.write(f"Заголовок {i}:\n")
                f.write(f"  Параграф: {header['para_idx']}\n")
                f.write(f"  Строк: {header['lines']}\n")
                f.write(f"  Имеет продолжение: {'Да' if header['has_continuation'] else 'Нет'}\n")
                f.write("  Текст:\n")

                # Show text with line breaks
                for line_num, line in enumerate(header['text'].split('\n'), 1):
                    f.write(f"    {line_num}: {line}\n")

                # Suggest XML output
                title_text = header['text'].split('\n')[0] if '\n' in header['text'] else header['text']
                f.write("  Предлагаемый XML:\n")
                f.write(f"    <levelledPara><title>{title_text}</title></levelledPara>\n")

                f.write("\n" + "-" * 60 + "\n\n")

        print(f"Analysis completed. Log saved to: {output_log}")

    except Exception as e:
        error_msg = f"Error analyzing document: {str(e)}"
        print(error_msg)
        with open(output_log, 'w', encoding='utf-8') as f:
            f.write(f"ERROR: {error_msg}\n")
        return

if __name__ == "__main__":
    # Document path
    doc_path = "docs/РСУ_адаптированная.docx"

    # Output log path
    output_log = "numbered_headers_analysis_log.txt"

    analyze_numbered_headers(doc_path, output_log)
