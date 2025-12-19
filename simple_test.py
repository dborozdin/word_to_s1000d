#!/usr/bin/env python3
"""
Простой тест для проверки многолистовых иллюстраций.
"""

import sys
import os
import tempfile
from docx import Document

# Добавляем путь к модулям проекта
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from parsers.multi_sheet_illustration_parser import find_multi_sheet_references
from processing_scripts.descriptive_processor import process_descriptive_document

def test_detection():
    """Тест обнаружения многолистовых иллюстраций."""
    print("=== ТЕСТ ОБНАРУЖЕНИЯ ===")
    
    test_elements = [
        {'type': 'paragraph', 'content': 'Рисунок 1 - Обычная иллюстрация'},
        {'type': 'paragraph', 'content': 'Рисунок 2 - Схема (лист 1, лист 2)'},
        {'type': 'paragraph', 'content': 'Рисунок 3 - Алгоритм (лист 1, лист 2, лист 3)'},
    ]
    
    multi_sheet_refs = find_multi_sheet_references(test_elements)
    print(f"Найдено: {len(multi_sheet_refs)} многолистовых иллюстраций")
    
    for ref in multi_sheet_refs:
        print(f"  Рисунок {ref['figure_number']}: листы {ref['sheets']}")
    
    return len(multi_sheet_refs) == 2

def test_processing():
    """Тест полной обработки."""
    print("\n=== ТЕСТ ОБРАБОТКИ ===")
    
    # Создаем простой документ
    doc = Document()
    doc.add_heading('Тест', 0)
    doc.add_paragraph('Рисунок 1 - Обычная иллюстрация')
    doc.add_paragraph('Рисунок 2 - Многолистовая (лист 1, лист 2)')
    
    # Сохраняем во временный файл
    temp_dir = tempfile.mkdtemp()
    test_file = os.path.join(temp_dir, 'test.docx')
    doc.save(test_file)
    
    # Обрабатываем
    output_dir = tempfile.mkdtemp()
    try:
        files = process_descriptive_document(test_file, output_dir)
        print(f"Создано файлов: {len(files)}")
        
        # Анализируем XML
        for xml_file in files:
            if xml_file.endswith('.xml'):
                with open(xml_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # Ищем figure элементы
                import re
                figures = re.findall(r'<figure[^>]*>.*?<title>(.*?)</title>(.*?)</figure>', 
                                   content, re.DOTALL | re.IGNORECASE)
                
                multi_sheet_count = 0
                for title, fig_content in figures:
                    graphic_count = len(re.findall(r'<graphic[^>]*>', fig_content))
                    if graphic_count > 1:
                        multi_sheet_count += 1
                        print(f"Многолистовая: {title.strip()}, листов: {graphic_count}")
                    else:
                        print(f"Обычная: {title.strip()}")
                
                return multi_sheet_count >= 1
                
    except Exception as e:
        print(f"Ошибка: {e}")
        return False
    finally:
        import shutil
        shutil.rmtree(temp_dir)
        shutil.rmtree(output_dir)

if __name__ == "__main__":
    print("🚀 Тестирование многолистовых иллюстраций")
    
    success1 = test_detection()
    success2 = test_processing()
    
    if success1 and success2:
        print("\n✅ ВСЕ ТЕСТЫ ПРОЙДЕНЫ!")
    else:
        print(f"\n❌ Тесты не пройдены: обнаружение={success1}, обработка={success2}")
