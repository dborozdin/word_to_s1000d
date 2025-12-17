#!/usr/bin/env python3
"""
Debug script to check specific paragraphs that should be numbered headers
"""

import os
import zipfile
import xml.etree.ElementTree as ET
from docx import Document

def debug_specific_paragraphs():
    doc = Document('docs/РСУ_адаптированная.docx')

    # Check paragraphs 1 and 11 specifically
    target_paragraphs = [1, 11]

    for para_idx in target_paragraphs:
        if para_idx < len(doc.paragraphs):
            paragraph = doc.paragraphs[para_idx]
            text = paragraph.text.strip()
            style_name = paragraph.style.name

            # Check for numbering properties
            numPr = None
            if hasattr(paragraph.paragraph_format, 'numPr') and paragraph.paragraph_format.numPr is not None:
                numPr = paragraph.paragraph_format.numPr

            print(f"\nParagraph {para_idx}:")
            print(f"  Text: '{text[:100]}...'")  # Truncate for readability
            print(f"  Style: {style_name}")
            print(f"  Has numbering: {numPr is not None}")

            if numPr is not None:
                if hasattr(numPr, 'numId') and numPr.numId is not None:
                    print(f"  numId: {numPr.numId}")
                if hasattr(numPr, 'ilvl') and numPr.ilvl is not None:
                    print(f"  ilvl: {numPr.ilvl}")

    # Now check the raw OOXML
    print("\n" + "="*50)
    print("RAW OOXML ANALYSIS")
    print("="*50)

    with zipfile.ZipFile('docs/РСУ_адаптированная.docx', 'r') as zip_ref:
        document_xml = zip_ref.read('word/document.xml')

    root = ET.fromstring(document_xml)
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    paragraphs = root.findall('.//w:p', ns)

    for i, para in enumerate(paragraphs):
        if i in target_paragraphs:
            text_runs = para.findall('.//w:t', ns)
            text = ''.join([run.text for run in text_runs if run.text])

            # Check for numPr in raw XML
            numPr_elem = para.find('.//w:numPr', ns)
            has_numPr_raw = numPr_elem is not None

            print(f"\nParagraph {i} (raw XML):")
            print(f"  Has numPr in raw XML: {has_numPr_raw}")

            if has_numPr_raw:
                ilvl_elem = numPr_elem.find('.//w:ilvl', ns)
                numId_elem = numPr_elem.find('.//w:numId', ns)
                if ilvl_elem is not None:
                    print(f"  Raw ilvl: {ilvl_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')}")
                if numId_elem is not None:
                    print(f"  Raw numId: {numId_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')}")

            # Show raw pPr section
            pPr = para.find('.//w:pPr', ns)
            if pPr is not None:
                print(f"  Raw pPr XML: {ET.tostring(pPr, encoding='unicode')[:300]}...")

if __name__ == "__main__":
    debug_specific_paragraphs()
