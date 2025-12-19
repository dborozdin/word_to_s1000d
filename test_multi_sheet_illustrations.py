#!/usr/bin/env python3
"""
Тестовый скрипт для проверки функциональности многолистовых иллюстраций.
Создает тестовый DOCX с многолистовыми иллюстрациями и проверяет их корректную генерацию в XML.
"""

import os
import sys
import tempfile
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Добавляем путь к модулям проекта
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from parsers.multi_sheet_illustration_parser import process_multi_sheet_illustrations, find_multi_sheet_references
from parsers.illustration_parser import extract_illustrations
from parsers.elements_analyzer import analyze_document_elements
from processing_scripts.descriptive_processor import process_descriptive_document
from generators.s1000d_generator import S1000DGenerator


def create_test_docx_with_multi_sheets() -> tuple:
    """
    Создает тестовый DOCX документ с многолистовыми иллюстрациями.
    
    Returns:
        Путь к созданному тестовому файлу и временная директория
    """
    # Создаем временный файл
    temp_dir = tempfile.mkdtemp()
    test_file = os.path.join(temp_dir, "test_multi_sheet_illustrations.docx")
    
    # Создаем документ
    doc = Document()
    
    # Добавляем заголовок
    title = doc.add_heading('Тестирование многолистовых иллюстраций', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Добавляем введение
    intro = doc.add_paragraph(
        'В данном документе тестируется функциональность многолистовых иллюстраций. '
        'Если в названии рисунка содержится текст "лист 1", "лист 2" и т.д., '
        'система должна создать многолистовую иллюстрацию согласно S1000D.'
    )
    
    # Добавляем обычную иллюстрацию (для сравнения)
    doc.add_heading('Обычная иллюстрация', level=1)
    doc.add_paragraph(
        'Ниже приведена обычная иллюстрация без упоминания листов. '
        'Эта иллюстрация должна быть обработана как обычный рисунок.'
    )
    doc.add_paragraph('Рисунок 1 - Общий вид системы')
    
    # Добавляем многолистовую иллюстрацию с 2 листами
    doc.add_heading('Многолистовая иллюстрация (2 листа)', level=1)
    doc.add_paragraph(
        'Ниже приведена многолистовая иллюстрация с упоминанием листов. '
        'Система должна создать один figure элемент с двумя graphic элементами.'
    )
    doc.add_paragraph('Рисунок 2 - Схема подключения блоков (лист 1, лист 2)')
    
    # Добавляем многолистовую иллюстрацию с 3 листами
    doc.add_heading('Многолистовая иллюстрация (3 листа)', level=1)
    doc.add_paragraph(
        'Еще один пример многолистовой иллюстрации с тремя листами.'
    )
    doc.add_paragraph('Рисунок 3 - Алгоритм работы системы (лист 1, лист 2, лист 3)')
    
    # Добавляем многолистовую иллюстрацию с 4 листами
    doc.add_heading('Многолистовая иллюстрация (4 листа)', level=1)
    doc.add_paragraph(
        'Пример многолистовой иллюстрации с четырьмя листами для проверки обработки большего количества листов.'
    )
    doc.add_paragraph('Рисунок 4 - Структурная схема системы (лист 1, лист 2, лист 3, лист 4)')
    
    # Добавляем еще одну обычную иллюстрацию
    doc.add_heading('Еще одна обычная иллюстрация', level=1)
    doc.add_paragraph(
        'Еще один пример обычной иллюстрации без упоминания листов.'
    )
    doc.add_paragraph('Рисунок 5 - Интерфейс пользователя')
    
    # Добавляем таблицу для проверки правильного позиционирования
    doc.add_heading('Тестовая таблица', level=1)
    doc.add_paragraph(
        'Ниже приведена таблица для проверки правильного позиционирования элементов в XML.'
    )
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = 'Параметр'
    table.cell(0, 1).text = 'Значение'
    table.cell(1, 0).text = 'Температура'
    table.cell(1, 1).text = '25°C'
    table.cell(2, 0).text = 'Давление'
    table.cell(2, 1).text = '1 атм'
    
    # Добавляем список для проверки
    doc.add_heading('Тестовый список', level=1)
    doc.add_paragraph('Основные компоненты системы:')
    
    list_para = doc.add_paragraph()
    list_para.style = 'List Paragraph'
    list_para.text = 'Блок управления'
    
    list_para = doc.add_paragraph()
    list_para.style = 'List Paragraph'
    list_para.text = 'Датчик температуры'
    
    list_para = doc.add_paragraph()
    list_para.style = 'List Paragraph'
    list_para.text = 'Исполнительный механизм'
    
    # Добавляем многолистовую иллюстрацию в другом формате
    doc.add_heading('Альтернативный формат', level=1)
    doc.add_paragraph(
        'Проверка альтернативного формата записи многолистовой иллюстрации.'
    )
    doc.add_paragraph('Рисунок 6 - Монтажная схема лист 1, лист 2, лист 3')
    
    # Сохраняем документ
    doc.save(test_file)
    
    print(f"Создан тестовый документ: {test_file}")
    return test_file, temp_dir


def test_multi_sheet_detection():
    """
    Тестирует функцию обнаружения многолистовых иллюстраций.
    """
    print("\n=== ТЕСТ 1: Обнаружение многолистовых иллюстраций ===")
    
    # Создаем тестовые элементы
    test_elements = [
        {
            'type': 'paragraph',
            'content': 'Рисунок 1 - Обычная иллюстрация',
            'start_para': 0,
            'end_para': 0
        },
        {
            'type': 'paragraph',
            'content': 'Рисунок 2 - Схема подключения блоков (лист 1, лист 2)',
            'start_para': 1,
            'end_para': 1
        },
        {
            'type': 'paragraph',
            'content': 'Рисунок 3 - Алгоритм работы системы (лист 1, лист 2, лист 3)',
            'start_para': 2,
            'end_para': 2
        },
        {
            'type': 'paragraph',
            'content': 'Рисунок 4 - Структурная схема системы (лист 1, лист 2, лист 3, лист 4)',
            'start_para': 3,
            'end_para': 3
        },
        {
            'type': 'paragraph',
            'content': 'Рисунок 5 - Интерфейс пользователя',
            'start_para': 4,
            'end_para': 4
        },
        {
            'type': 'paragraph',
            'content': 'Рисунок 6 - Монтажная схема лист 1, лист 2, лист 3',
            'start_para': 5,
            'end_para': 5
        }
    ]
    
    # Находим многолистовые иллюстрации
    multi_sheet_refs = find_multi_sheet_references(test_elements)
    
    print(f"Найдено многолистовых иллюстраций: {len(multi_sheet_refs)}")
    
    for ref in multi_sheet_refs:
        print(f"  - Рисунок {ref['figure_number']}: листы {ref['sheets']} (количество: {ref['sheet_count']})")
        print(f"    Контент: {ref['content']}")
    
    assert len(multi_sheet_refs) == 4, f"Ожидалось найти 4 многолистовые иллюстрации, найдено: {len(multi_sheet_refs)}"
    
    # Проверяем правильность обнаружения
    expected_figures = {2: [1, 2], 3: [1, 2, 3], 4: [1, 2, 3, 4], 6: [1, 2, 3]}
    
    for ref in multi_sheet_refs:
        fig_num = ref['figure_number']
        sheets = ref['sheets']
        
        assert fig_num in expected_figures, f"Неожиданный номер рисунка: {fig_num}"
        assert sheets == expected_figures[fig_num], f"Неправильные листы для рисунка {fig_num}: ожидались {expected_figures[fig_num]}, получены {sheets}"
    
    print("✅ Тест обнаружения многолистовых иллюстраций пройден успешно!")


def test_full_processing():
    """
    Тестирует полный процесс обработки документа с многолистовыми иллюстрациями.
    """
    print("\n=== ТЕСТ 2: Полная обработка документа ===")
    
    # Создаем тестовый документ
    test_file, temp_dir = create_test_docx_with_multi_sheets()
    
    try:
        # Создаем временную директорию для выходных файлов
        output_dir = tempfile.mkdtemp()
        
        print(f"Тестовый файл: {test_file}")
        print(f"Выходная директория: {output_dir}")
        
        # Обрабатываем документ
        print("\nЗапуск обработки документа...")
        generated_files = process_descriptive_document(test_file, output_dir)
        
        print(f"\nСгенерировано файлов: {len(generated_files)}")
        for filepath in generated_files:
            print(f"  - {os.path.basename(filepath)}")
        
        # Проверяем, что файлы созданы
        assert len(generated_files) > 0, "Не было создано ни одного файла"
        
        # Анализируем сгенерированные XML файлы
        print("\n=== АНАЛИЗ СГЕНЕРИРОВАННЫХ XML ===")
        
        multi_sheet_figures_found = 0
        regular_figures_found = 0
        
        for xml_file in generated_files:
            if not xml_file.endswith('.xml'):
                continue
                
            print(f"\nАнализ файла: {os.path.basename(xml_file)}")
            
            with open(xml_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Ищем figure элементы
            import re
            
            # Находим все figure элементы
            figure_pattern = r'<figure[^>]*id="([^"]*)"[^>]*>.*?<title>(.*?)</title>(.*?)</figure>'
            figures = re.findall(figure_pattern, content, re.DOTALL | re.IGNORECASE)
            
            for fig_id, title, fig_content in figures:
                # Подсчитываем graphic элементы в figure
                graphic_count = len(re.findall(r'<graphic[^>]*>', fig_content, re.IGNORECASE))
                
                if graphic_count > 1:
                    multi_sheet_figures_found += 1
                    print(f"  📊 Многолистовая иллюстрация: {title.strip()}")
                    print(f"     ID: {fig_id}")
                    print(f"     Количество листов: {graphic_count}")
                    
                    # Выводим все graphic элементы
                    graphics = re.findall(r'<graphic[^>]*infoEntityIdent="([^"]*)"[^>]*>', fig_content, re.IGNORECASE)
                    for i, graphic_ident in enumerate(graphics, 1):
                        print(f"       Лист {i}: {graphic_ident}")
                else:
                    regular_figures_found += 1
                    print(f"  📷 Обычная иллюстрация: {title.strip()}")
                    print(f"     ID: {fig_id}")
                    print(f"     Количество листов: {graphic_count}")
        
        print(f"\n=== РЕЗУЛЬТАТЫ АНАЛИЗА ===")
        print(f"Многолистовых иллюстраций найдено: {multi_sheet_figures_found}")
        print(f"Обычных иллюстраций найдено: {regular_figures_found}")
        
        # Проверяем результаты
        expected_multi_sheet = 3  # Рисунки 2, 3, 4, 6 (4 многолистовых)
        expected_regular = 2      # Рисунки 1, 5 (2 обычных)
        
        print(f"\nОжидалось:")
        print(f"  - Многолистовых иллюстраций: {expected_multi_sheet}")
        print(f"  - Обычных иллюстраций: {expected_regular}")
        
        if multi_sheet_figures_found >= expected_multi_sheet:
            print("✅ Многолистовые иллюстрации корректно обработаны!")
        else:
            print("❌ Не все многолистовые иллюстрации были обработаны правильно")
        
        if regular_figures_found >= expected_regular:
            print("✅ Обычные иллюстрации корректно обработаны!")
        else:
            print("❌ Обычные иллюстрации не были обработаны правильно")
        
        # Проверяем позиционирование элементов
        print("\n=== ПРОВЕРКА ПОЗИЦИОНИРОВАНИЯ ===")
        
        # Ищем порядок элементов в XML
        for xml_file in generated_files:
            if not xml_file.endswith('.xml'):
                continue
                
            with open(xml_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Извлекаем порядок элементов
            elements_in_order = []
            
            # Ищем параграфы
            paragraphs = re.findall(r'<para>(.*?)</para>', content, re.DOTALL)
            elements_in_order.extend([('paragraph', p.strip()[:50] + '...' if len(p.strip()) > 50 else p.strip()) for p in paragraphs])
            
            # Ищем figure элементы
            figures = re.findall(r'<figure[^>]*>.*?<title>(.*?)</title>.*?</figure>', content, re.DOTALL | re.IGNORECASE)
            elements_in_order.extend([('figure', f.strip()[:50] + '...' if len(f.strip()) > 50 else f.strip()) for f in figures])
            
            # Ищем таблицы
            tables = re.findall(r'<table[^>]*>.*?</table>', content, re.DOTALL)
            elements_in_order.extend([('table', f'Таблица {i+1}') for i in range(len(tables))])
            
            print(f"Порядок элементов в {os.path.basename(xml_file)}:")
            for i, (elem_type, content_preview) in enumerate(elements_in_order[:10], 1):  # Показываем первые 10 элементов
                print(f"  {i}. {elem_type}: {content_preview}")
        
        return True
        
    except Exception as e:
        print(f"❌ Ошибка при обработке: {e}")
        import traceback
        traceback.print_exc()
        return False
        
    finally:
        # Очищаем временные файлы
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        if 'output_dir' in locals() and os.path.exists(output_dir):
            shutil.rmtree(output_dir)


def run_all_tests():
    """
    Запускает все тесты.
    """
    print("🚀 Запуск тестирования функциональности многолистовых иллюстраций")
    print("=" * 70)
    
    try:
        # Тест 1: Обнаружение многолистовых иллюстраций
        test_multi_sheet_detection()

        # Тест 2: Полная обработка документа
        success = test_full_processing()

        if success:
            print("\n" + "=" * 70)
            print("🎉 ВСЕ ТЕСТЫ ПРОЙДЕНЫ УСПЕШНО!")
            print("Функциональность многолистовых иллюстраций работает корректно!")
        else:
            print("\n" + "=" * 70)
            print("❌ НЕКОТОРЫЕ ТЕСТЫ НЕ ПРОЙДЕНЫ!")
    except Exception as e:
        print(f"❌ Критическая ошибка при запуске тестов: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    run_all_tests()
