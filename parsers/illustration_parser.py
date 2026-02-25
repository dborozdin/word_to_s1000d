"""
Illustration parser for docx documents.
Extracts embedded images and tracks their references.
"""

import os
import re
import shutil
from typing import Dict, List, Tuple
from docx import Document
from docx.shared import Inches
from PIL import Image as PILImage, ImageDraw, ImageFont
from docx.document import Document as DocxDocument
from docx.oxml.shape import CT_Picture


def extract_illustrations(doc: Document, output_dir: str = "./tg_web/publications", graphic_ident_prefix: str = None) -> Tuple[Dict[str, str], Dict[str, Dict]]:
    """
    Extract embedded images from document and save to output/graphics directory with S1000D naming.
    Images are numbered based on the order of illustration references in the text.

    Args:
        doc: Docx document object
        output_dir: Output directory (parent of graphics subfolder)

    Returns:
        Tuple of:
        - Dictionary mapping reference names to image file paths
        - Dictionary mapping reference names to position info for elements log
    """
    # Create graphics subdirectory in output
    graphics_dir = os.path.join(output_dir, "graphics")
    if not os.path.exists(graphics_dir):
        os.makedirs(graphics_dir)

    illustrations = {}
    illustration_positions = {}

    # Step 1: Scan text to find illustration references in order of appearance
    reference_order = _get_illustration_reference_order(doc)

    # Step 2: Create mapping from reference number to sequential GRAPHIC number
    reference_to_graphic = {}
    for idx, ref_num in enumerate(reference_order):
        reference_to_graphic[ref_num] = idx

    # Step 3: Extract embedded images in document order
    embedded_images = []

    # Track cumulative text for position calculation
    line_number = 1
    char_position = 0

    # Process document elements to find embedded images in order of appearance
    for i, (element_type, element_idx) in enumerate(_get_doc_elements(doc)):
        if element_type == 'paragraph':
            paragraph = doc.paragraphs[element_idx]

            # Update position tracking
            para_start_line = line_number
            para_start_char = char_position

            # Add paragraph text to cumulative text
            original_text = paragraph.text
            newline_count = original_text.count('\n')
            line_number += newline_count
            if newline_count > 0:
                char_position = len(original_text.split('\n')[-1])
            else:
                char_position += len(original_text)

            para_end_line = line_number
            para_end_char = char_position

            # Check if paragraph contains embedded image
            if _has_embedded_image(paragraph):
                try:
                    # Get the specific rId for this image
                    rId = _get_image_rid_from_paragraph(paragraph)
                    if rId:
                        # Find the corresponding relationship and extract image
                        rel = doc.part.rels.get(rId)
                        if rel and "image" in rel.target_ref:
                            img_blob = rel.target_part.blob
                            embedded_images.append({
                                'blob': img_blob,
                                'start_line': para_start_line,
                                'start_char': para_start_char,
                                'end_line': para_end_line,
                                'end_char': para_end_char,
                                'start_para': element_idx,
                                'end_para': element_idx,
                                'context_text': original_text.strip()[:100] + ('...' if len(original_text.strip()) > 100 else '')
                            })

                except Exception as e:
                    print(f"Error extracting image: {e}")

        elif element_type == 'table':
            table = doc.tables[element_idx]

            # Update position tracking for table
            table_start_line = line_number
            table_start_char = char_position

            # Process each cell in the table
            for row in table.rows:
                for cell in row.cells:
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        # Update position for each paragraph in cell
                        cell_para_start_line = line_number
                        cell_para_start_char = char_position

                        # Add paragraph text to cumulative text
                        original_text = paragraph.text
                        newline_count = original_text.count('\n')
                        line_number += newline_count
                        if newline_count > 0:
                            char_position = len(original_text.split('\n')[-1])
                        else:
                            char_position += len(original_text)

                        cell_para_end_line = line_number
                        cell_para_end_char = char_position

                        # Check if paragraph in table cell contains embedded image
                        if _has_embedded_image(paragraph):
                            try:
                                # Get the specific rId for this image
                                rId = _get_image_rid_from_paragraph(paragraph)
                                if rId:
                                    # Find the corresponding relationship and extract image
                                    rel = doc.part.rels.get(rId)
                                    if rel and "image" in rel.target_ref:
                                        img_blob = rel.target_part.blob
                                        embedded_images.append({
                                            'blob': img_blob,
                                            'start_line': cell_para_start_line,
                                            'start_char': cell_para_start_char,
                                            'end_line': cell_para_end_line,
                                            'end_char': cell_para_end_char,
                                            'start_para': element_idx,  # table index
                                            'end_para': element_idx,
                                            'context_text': original_text.strip()[:100] + ('...' if len(original_text.strip()) > 100 else '')
                                        })

                            except Exception as e:
                                print(f"Error extracting image from table: {e}")

            table_end_line = line_number
            table_end_char = char_position

    # Create log file for illustration extraction
    _logs_dir = os.path.join(output_dir, '_logs')
    os.makedirs(_logs_dir, exist_ok=True)
    log_path = os.path.join(_logs_dir, "illustration_extraction.log")
    with open(log_path, 'w', encoding='utf-8') as log_file:
        log_file.write("Лог извлечения иллюстраций\n")
        log_file.write("=" * 50 + "\n")
        log_file.write("Индекс | Имя файла | Строка | Символ\n")
        log_file.write("-" * 50 + "\n")

    # Step 4: Save images with sequential GRAPHIC numbering in document order
    effective_prefix = graphic_ident_prefix or "GS5-A-120-10-00-00A-041A-A_001_RU-RU"
    for idx, img_info in enumerate(embedded_images):
        # Use sequential numbering starting from 0 for all extracted images
        graphic_num = idx

        img_name = f"{effective_prefix}-GRAPHIC{graphic_num}.jpg"
        img_path = os.path.join(graphics_dir, img_name)

        # Save image
        with open(img_path, 'wb') as f:
            f.write(img_info['blob'])

        # Map to reference name
        ref_name = f"{effective_prefix}-GRAPHIC{graphic_num}"
        illustrations[ref_name] = img_path

        # Store position information
        illustration_positions[ref_name] = {
            'start_line': img_info['start_line'],
            'start_char': img_info['start_char'],
            'end_line': img_info['end_line'],
            'end_char': img_info['end_char'],
            'start_para': img_info['start_para'],
            'end_para': img_info['end_para'],
            'context_text': img_info['context_text']
        }

        # Write to log file
        with open(log_path, 'a', encoding='utf-8') as log_file:
            log_file.write(f"{graphic_num} | {img_name} | {img_info['start_line']} | {img_info['start_char']}\n")

    return illustrations, illustration_positions


def _get_illustration_reference_order(doc: Document) -> List[int]:
    """
    Scan document text to find illustration references in order of appearance.

    Returns:
        List of unique reference numbers in order of first appearance
    """
    seen_refs = set()
    reference_order = []

    for paragraph in doc.paragraphs:
        text = paragraph.text
        # Find illustration references
        patterns = [
            r'\b[Рр]исунок\s*(\d+)',
            r'\b[Рр]ис\.\s*(\d+)',
            r'\b[Фф]igure\s*(\d+)',
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                ref_num = int(match)
                if ref_num not in seen_refs:
                    seen_refs.add(ref_num)
                    reference_order.append(ref_num)

    return reference_order


def _get_doc_elements(doc: Document) -> List[Tuple[str, int]]:
    """Get all document elements (paragraphs and tables) in order."""
    elements = []
    para_idx = 0
    table_idx = 0

    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            if para_idx < len(doc.paragraphs):
                elements.append(('paragraph', para_idx))
                para_idx += 1
        elif element.tag.endswith('tbl'):  # Table
            if table_idx < len(doc.tables):
                elements.append(('table', table_idx))
                table_idx += 1

    return elements


def _has_embedded_image(paragraph) -> bool:
    """Check if paragraph contains embedded image (DrawingML or VML)."""
    for run in paragraph.runs:
        # DrawingML format (modern .docx)
        if run.element.xpath('.//a:blip'):
            return True
    # VML format (converted .doc → .docx via Word COM)
    # v:imagedata lives inside w:pict which may be outside runs
    para_xml = paragraph._element.xml
    if 'v:imagedata' in para_xml:
        return True
    return False


def _get_image_rid_from_paragraph(paragraph) -> str:
    """Extract relationship ID (rId) from embedded image in paragraph."""
    # Try DrawingML first (a:blip)
    for run in paragraph.runs:
        blip_elements = run.element.xpath('.//a:blip')
        if blip_elements:
            blip = blip_elements[0]
            embed_attr = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if embed_attr:
                return embed_attr
    # Try VML format (v:imagedata inside w:pict)
    VML_NS = 'urn:schemas-microsoft-com:vml'
    R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    for elem in paragraph._element.iter('{%s}imagedata' % VML_NS):
        rid = elem.get('{%s}id' % R_NS)
        if rid:
            return rid
    return None


def find_image_references(text: str) -> List[Tuple[str, str]]:
    """
    Find image references in text.

    Args:
        text: Text content to search

    Returns:
        List of (reference_type, reference_number) tuples
    """
    references = []

    # Look for figure references like "рисунок 1", "figure 1", etc.
    figure_patterns = [
        r'[Рр]исунок\s*(\d+)',
        r'[Рр]ис\.\s*(\d+)',
        r'[Фф]igure\s*(\d+)'
    ]

    for pattern in figure_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            references.append(('figure', match))

    return references


def map_figures_to_illustrations(figure_refs: List[Tuple[str, str]], illustrations: Dict[str, str], graphic_ident_prefix: str = None) -> Dict[str, str]:
    """
    Map figure references to actual illustration files.

    Args:
        figure_refs: List of (type, number) references
        illustrations: Dict of illustration reference to file path
        graphic_ident_prefix: Optional prefix for graphic naming (default: hardcoded RSUO prefix)

    Returns:
        Dictionary mapping figure references to illustration files
    """
    mapping = {}
    effective_prefix = graphic_ident_prefix or "GS5-A-120-10-00-00A-041A-A_001_RU-RU"

    for ref_type, ref_num in figure_refs:
        # Map to S1000D GRAPHIC naming: GRAPHIC{N} where N starts from 0
        # Assume figure reference numbers are 1-indexed, convert to 0-indexed
        graphic_num = int(ref_num) - 1
        graphic_name = f"{effective_prefix}-GRAPHIC{graphic_num}"

        if graphic_name in illustrations:
            mapping[f"{ref_type}_{ref_num}"] = illustrations[graphic_name]
        else:
            # Fallback: try exact mapping or other variations
            print(f"Warning: Could not find illustration for {ref_type} {ref_num}")

    return mapping


def get_document_illustrations_info(doc: Document) -> Dict[str, str]:
    """
    Get basic information about illustrations in the document.

    Args:
        doc: Docx document object

    Returns:
        Dictionary with illustration count and paths
    """
    total_images = 0

    # Count inline shapes with images
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.element.xpath('.//a:blip'):
                total_images += 1

    # Count shapes in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.element.xpath('.//a:blip'):
                            total_images += 1

    return {
        'total_count': total_images,
        'message': f'Document contains approximately {total_images} embedded images'
    }


def generate_placeholder_image(filepath: str, text: str = "< Добавьте иллюстрацию >"):
    """
    Generate a placeholder JPEG image with the given text.

    Args:
        filepath: Full path where the placeholder .jpg should be saved
        text: Text to render on the placeholder image
    """
    width, height = 800, 200
    img = PILImage.new('RGB', (width, height), color=(240, 240, 240))
    draw = ImageDraw.Draw(img)

    # Try to use a reasonable font size; fall back to default
    try:
        font = ImageFont.truetype("arial.ttf", 28)
    except (OSError, IOError):
        font = ImageFont.load_default()

    bbox = draw.textbbox((0, 0), text, font=font)
    text_w = bbox[2] - bbox[0]
    text_h = bbox[3] - bbox[1]
    x = (width - text_w) // 2
    y = (height - text_h) // 2
    draw.text((x, y), text, fill=(180, 0, 0), font=font)

    # Draw border
    draw.rectangle([0, 0, width - 1, height - 1], outline=(200, 200, 200), width=2)

    img.save(filepath, 'JPEG', quality=85)


def ensure_missing_placeholders(figure_info: List[Dict], output_dir: str):
    """
    Check figure_info for referenced graphic files that don't exist
    on disk and generate placeholder images for them.

    Args:
        figure_info: List of dicts with 'file' key (e.g. 'PREFIX-GRAPHIC0.jpg')
        output_dir: Output directory (parent of graphics/)
    """
    if not figure_info:
        return

    graphics_dir = os.path.join(output_dir, "graphics")
    if not os.path.exists(graphics_dir):
        os.makedirs(graphics_dir)

    for fig in figure_info:
        filename = fig.get('file', '')
        if not filename:
            continue
        filepath = os.path.join(graphics_dir, filename)
        if not os.path.exists(filepath):
            print(f"  Generating placeholder for missing illustration: {filename}")
            generate_placeholder_image(filepath)


def copy_publication_logo(output_dir: str):
    """
    Copy publication_logo.JPG to the output graphics/ directory
    so the viewer can resolve the PUBLICATION_LOGO entity.

    Args:
        output_dir: Output directory (parent of graphics/)
    """
    graphics_dir = os.path.join(output_dir, "graphics")
    dest = os.path.join(graphics_dir, "publication_logo.JPG")
    if os.path.exists(dest):
        return

    # Search for source logo in known locations
    script_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    candidates = [
        os.path.join(script_dir, "manual_data_modules", "graphics", "publication_logo.JPG"),
    ]

    for src in candidates:
        if os.path.exists(src):
            if not os.path.exists(graphics_dir):
                os.makedirs(graphics_dir)
            shutil.copy2(src, dest)
            return
