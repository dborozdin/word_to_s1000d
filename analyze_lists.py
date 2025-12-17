#!/usr/bin/env python3
"""
Script to analyze numbered lists in the test document РСУ_адаптированная.docx
"""

import os
import sys
from docx import Document
from parsers.list_parser import extract_lists

def analyze_lists_in_document(doc_path: str, output_log: str):
    """
    Analyze lists in document and write results to log file.

    Args:
        doc_path: Path to the document
        output_log: Path for output log file
    """
    print(f"Analyzing lists in document: {doc_path}")

    # Check if document exists
    if not os.path.exists(doc_path):
        error_msg = f"Document not found: {doc_path}"
        print(error_msg)
        with open(output_log, 'w', encoding='utf-8') as f:
            f.write(f"ERROR: {error_msg}\n")
        return

    try:
        # Load document
        doc = Document(doc_path)
        print(f"Document loaded successfully. Total paragraphs: {len(doc.paragraphs)}")

        # Extract lists
        lists_data = extract_lists(doc)
        print(f"Found {len(lists_data)} lists")

        # Write log
        with open(output_log, 'w', encoding='utf-8') as f:
            f.write("Анализ нумерованных списков в документе\n")
            f.write("=" * 50 + "\n")
            f.write(f"Документ: {os.path.basename(doc_path)}\n")
            f.write(f"Всего параграфов: {len(doc.paragraphs)}\n")
            f.write(f"Найдено списков: {len(lists_data)}\n\n")

            if not lists_data:
                f.write("Нумерованные списки не найдены в документе.\n")
                return

            for i, list_info in enumerate(lists_data, 1):
                f.write(f"Список {i}:\n")
                f.write(f"  Тип: {list_info.get('type', 'неизвестный')}\n")
                f.write(f"  Количество элементов: {len(list_info.get('items', []))}\n")
                f.write("\n  Элементы:\n")

                for j, item in enumerate(list_info.get('items', []), 1):
                    f.write(f"    {j}. {item}\n")

                f.write("\n" + "-" * 50 + "\n\n")

            # Summary
            numbered_count = sum(1 for lst in lists_data if lst.get('type') == 'numbered_list')
            unnumbered_count = sum(1 for lst in lists_data if lst.get('type') == 'unnumbered_list')

            f.write("Сводка:\n")
            f.write(f"  Нумерованных списков: {numbered_count}\n")
            f.write(f"  Маркированных списков: {unnumbered_count}\n")
            f.write(f"  Всего списков: {len(lists_data)}\n")

        print(f"Analysis completed. Log saved to: {output_log}")

    except Exception as e:
        error_msg = f"Error analyzing document: {str(e)}"
        print(error_msg)
        with open(output_log, 'w', encoding='utf-8') as f:
            f.write(f"ERROR: {error_msg}\n")
        return

if __name__ == "__main__":
    # Document path
    doc_path = "docs/РСУ_адаптированная.docx"

    # Output log path
    output_log = "list_analysis_log.txt"

    analyze_lists_in_document(doc_path, output_log)
