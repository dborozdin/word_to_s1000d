"""
Descriptive module processor.
Orchestrate parsing, maps sections to infoCodes, and generates multiple XML files.
"""

import os
import re
import configparser
from typing import Dict, List
from docx import Document

# Import parser modules
from parsers.text_parser import extract_text_by_headings, get_document_structure
from parsers.table_parser import get_tables_by_reference
from parsers.list_parser import extract_lists, convert_list_to_s1000d_randomlist
from parsers.illustration_parser import extract_illustrations, find_image_references, map_figures_to_illustrations
from parsers.content_analyzer import analyze_document_content, generate_content_analysis_log, get_content_for_info_name, generate_module_mapping_log
from parsers.elements_analyzer import analyze_document_elements, generate_elements_log

# Import generators
from generators.s1000d_generator import S1000DGenerator, create_data_module_config
from generators.pm_generator import PMGenerator, create_pm_config, create_dm_ref_data


def normalize_title_case(text: str) -> str:
    """
    Normalize title case for techName according to requirements.

    Args:
        text: Input text

    Returns:
        Normalized text
    """
    if not text:
        return text
    
    # If all letters are uppercase, leave as is
    if text.isupper():
        return text
    
    # Split by parentheses, preserving the parentheses
    parts = re.split(r'(\([^)]*\))', text)
    
    result_parts = []
    for part in parts:
        if part.startswith('(') and part.endswith(')'):
            # Keep parentheses content as is
            result_parts.append(part)
        else:
            # Apply title case: first letter uppercase, rest lowercase
            result_parts.append(part.capitalize())
    
    return ''.join(result_parts)


def extract_document_title(document: Document) -> str:
    """
    Extract main title from the first page/document header.

    Args:
        document: Word document object

    Returns:
        Document title or default if not found
    """
    # First try to get from the first paragraph (document title often has Normal style)
    for para in document.paragraphs:
        text = para.text.strip()
        if text and not para.style.name.startswith('Heading'):  # Skip headings, look for title
            # Check if it looks like a title with dash
            if "–" in text:
                title = text.split("–")[0].strip()
                return normalize_title_case(title)
            elif "-" in text:
                title = text.split("-")[0].strip()
                return normalize_title_case(title)
            # If it has multiple words and is long, might be title
            if len(text) > 20:
                return normalize_title_case(text.strip())

    # Fallback: get from headings if above didn't work
    headings = get_document_structure(document)
    if headings:
        first_heading = headings[0].strip()
        # Parse the part before "–" or "-"
        if "–" in first_heading:
            title = first_heading.split("–")[0].strip()
            return normalize_title_case(title)
        elif "-" in first_heading:
            title = first_heading.split("-")[0].strip()
            return normalize_title_case(title)

    # Ultimate fallback
    title = "Система"

    # Apply title case normalization
    return normalize_title_case(title)


def extract_organization_from_document(document: Document) -> str:
    """
    Extract organization name from document headers and footers.

    Args:
        document: Word document object

    Returns:
        Organization name or "Организация не задана" if not found
    """
    text = ""

    # Check all sections' headers and footers
    for section in document.sections:
        # Get header text - headers can have different types (first_page, even_page, odd_page)
        if hasattr(section.header, 'is_linked_to_previous') and not section.header.is_linked_to_previous:
            # Only add if not linked to avoid duplication
            for paragraph in section.header.paragraphs:
                text += paragraph.text + " "

        # Get footer text
        if hasattr(section.footer, 'is_linked_to_previous') and not section.footer.is_linked_to_previous:
            for paragraph in section.footer.paragraphs:
                text += paragraph.text + " "

    # Regex to find organization: abbreviation + nearest word after it
    pattern = r'(?:ОКБ|КБ|АО|ООО|НИЦ)\s*([^\s]+)'
    match = re.search(pattern, text, re.IGNORECASE)

    if match:
        abbreviation = match.group(0).replace(match.group(1), '').strip()
        word_after = match.group(1).strip()
        return f"{abbreviation} {word_after}"
    else:
        return "Организация не задана"


def map_heading_to_info_code(heading: str, component_index: int = 0) -> Dict:
    """
    Map document heading to S1000D infoCode.

    Args:
        heading: Document heading text
        component_index: Index for component numbering

    Returns:
        Dict with DM code components
    """
    heading_lower = heading.lower()

    base_dm_code = {
        'modelIdentCode': 'S5',
        'systemDiffCode': 'A',
        'systemCode': '120',
        'subSystemCode': '1',  # RSOO subsystem
        'subSubSystemCode': '0',
        'assyCode': '00',
        'disassyCode': '00',
        'disassyCodeVariant': 'A',
        'infoCode': '012',
        'infoCodeVariant': 'A',
        'itemLocationCode': 'A'
    }

    # System level modules
    if 'общие сведения' in heading_lower:
        base_dm_code.update({'infoCode': '011', 'infoCodeVariant': 'A'})  # Purpose
        return base_dm_code
    elif 'состав' in heading_lower or 'описание' in heading_lower:
        base_dm_code.update({'infoCode': '012', 'infoCodeVariant': 'A'})  # Description
        return base_dm_code
    elif 'структурно представляет' in heading_lower:
        base_dm_code.update({'infoCode': '013', 'infoCodeVariant': 'A'})  # Operation structure
        return base_dm_code
    elif 'информацион' in heading_lower:
        base_dm_code.update({'infoCode': '014', 'infoCodeVariant': 'A'})  # Info exchange
        return base_dm_code
    elif 'режимы работы' in heading_lower:
        base_dm_code.update({'infoCode': '013', 'infoCodeVariant': 'A'})  # General operation
        return base_dm_code
    elif 'автомати' in heading_lower in heading_lower:
        base_dm_code.update({'infoCode': '015', 'infoCodeVariant': 'A'})  # Automatic mode
        return base_dm_code
    elif 'автоном' in heading_lower in heading_lower:
        base_dm_code.update({'infoCode': '015', 'infoCodeVariant': 'B'})  # Autonomous mode
        return base_dm_code
    elif 'аварийн' in heading_lower in heading_lower:
        base_dm_code.update({'infoCode': '015', 'infoCodeVariant': 'C'})  # Emergency mode
        return base_dm_code
    elif 'учебно' in heading_lower in heading_lower:
        base_dm_code.update({'infoCode': '015', 'infoCodeVariant': 'D'})  # Training mode
        return base_dm_code
    elif 'управлени' in heading_lower and ('створками' in heading_lower or 'платформ' in heading_lower):
        base_dm_code.update({'infoCode': '016', 'infoCodeVariant': 'A'})  # Additional operations
        return base_dm_code
    else:
        # Check if it's a component (equipment, unit, block, etc.)
        unit_keywords = ['блок', 'пульт', 'рама', 'система', 'блокировка', 'контроллер', 'бусп', 'ски', 'бкп', 'пусть', 'пнп', 'би-м', 'сброса', 'псо-5п', 'бус-5пм', 'рм-5п', 'рм-5пм', 'рм-5п-1', 'пусть-5пм', 'бч-70']
        if any(keyword in heading_lower for keyword in unit_keywords):
            # Component description
            base_dm_code.update({
                'subSubSystemCode': f'{component_index % 10}',
                'infoCode': '017',
                'infoCodeVariant': 'A'
            })
            return base_dm_code

    # Default to description
    base_dm_code.update({'infoCode': '012'})
    return base_dm_code


def process_descriptive_document(doc_path: str, output_dir: str):
    """
    Process descriptive document: orchestrate parsing, map sections, generate XML files.

    Args:
        doc_path: Path to docx document
        output_dir: Output directory for generated files
    """
    print(f"Processing descriptive document: {doc_path}")

    # Read configuration
    config = configparser.ConfigParser()
    config.read('config.ini')
    split_into_modules = config.getboolean('processing', 'split_into_modules', fallback=False)
    print(f"Split into modules: {split_into_modules}")

    # Load document
    doc = Document(doc_path)

    # Extract organization from headers/footers
    organization = extract_organization_from_document(doc)
    print(f"Extracted organization: {organization}")

    # Extract document title
    document_title = extract_document_title(doc)
    print(f"Extracted document title: {document_title}")

    # Parse illustrations first to get correct infoEntityIdent mappings
    print("Extracting illustrations...")
    illustrations, illustration_positions = extract_illustrations(doc, output_dir)

    # Analyze document content using the new content analyzer
    print("Analyzing document content structure...")
    analysis_results = analyze_document_content(doc)

    # Generate content analysis log (only if splitting into modules)
    log_path = None
    if split_into_modules:
        log_path = generate_content_analysis_log(doc_path, analysis_results, output_dir)

    # Analyze document elements for logging and XML generation
    print("Analyzing document elements...")
    elements = analyze_document_elements(doc, illustrations, illustration_positions)

    # Generate elements log (always)
    elements_log_path = generate_elements_log(doc_path, elements, output_dir)

    # Parse additional content types
    print("Parsing additional content...")
    text_sections = extract_text_by_headings(doc)
    tables = get_tables_by_reference(doc)
    lists_data = extract_lists(doc)

    print(f"Found {len(analysis_results)} analyzed sections, {len(tables)} tables, {len(lists_data)} lists, {len(illustrations)} illustrations")

    # Initialize generator
    generator = S1000DGenerator()
    pm_generator = PMGenerator()

    # Generate data modules based on content analysis
    generated_files = []
    dm_refs = []  # List to collect DM references for PM
    component_counter = 0

    # Track mapping of sections to modules for logging
    module_mapping = {}  # filename -> module_data

    # Group sections by their unique characteristics for module generation
    section_groups = group_sections_for_modules(analysis_results, split_into_modules)

    for group_key, group_info in section_groups.items():
        sections_in_group = group_info['sections']

        # Choose representative section for DM code generation (first one)
        representative_section = sections_in_group[0]

        # Prepare combined content for this module
        combined_content = {
            "xml_parts": []
        }

        # Track global counters for unique IDs across all sections
        illustration_counter = 0
        figure_counter = 0
        figure_info = []

        if not split_into_modules:
            # For combined case, create one big section covering all paragraphs
            big_section = {
                'start_para': 0,
                'end_para': len(doc.paragraphs) - 1,
                'info_name': group_info['info_name']
            }
            section_content = assemble_content_for_section(big_section, doc, tables, lists_data, elements, illustration_counter, figure_counter, figure_info, document_title, group_info['info_name'])
            combined_content["xml_parts"].extend(section_content["xml_parts"])
        else:
            # Combine content from all sections in this group
            for section in sections_in_group:
                section_content = assemble_content_for_section(section, doc, tables, lists_data, elements, illustration_counter, figure_counter, figure_info, document_title, section.get('info_name', ''))
                combined_content["xml_parts"].extend(section_content["xml_parts"])
                # Update counters from the section processing
                illustration_counter = section_content.get("illustration_counter", illustration_counter)
                figure_counter = section_content.get("figure_counter", figure_counter)

        # Determine DM code based on section type
        dm_code = get_dm_code_for_section(representative_section, component_counter)

        # Map section type to appropriate module type
        if representative_section.get('section_type') == 'purpose':
            if representative_section.get('info_name') == 'Описание функций изделия':
                dm_code.update({'infoCode': '012', 'infoCodeVariant': 'B'})  # Function description
            else:
                dm_code.update({'infoCode': '011', 'infoCodeVariant': 'A'})  # General purpose
        elif representative_section.get('section_type') == 'description':
            dm_code.update({'infoCode': '012', 'infoCodeVariant': 'A'})  # Description
        elif representative_section.get('section_type') == 'operation':
            dm_code.update({'infoCode': '013', 'infoCodeVariant': 'A'})  # Operation
        elif representative_section.get('section_type') == 'component':
            dm_code.update({'infoCode': '017', 'infoCodeVariant': 'A'})  # Component description
            component_counter += 1

        # Create DM config
        dm_config = create_data_module_config(
            document_title,
            representative_section.get('info_name', 'Неопределен'),
            dm_code,
            combined_content,
            enterprise_name=organization,
            originator_name=organization
        )

        # Generate XML file
        filepath = generator.generate_data_module(dm_config, output_dir, illustrations, figure_info)
        generated_files.append(filepath)

        # Collect DM ref for PM generation
        dm_ref = create_dm_ref_data(
            dm_code,
            document_title,  # techName
            representative_section.get('info_name', 'Неопределен')  # infoName
        )
        dm_refs.append(dm_ref)

        # Track mapping for logging
        filename = os.path.basename(filepath)
        module_mapping[filename] = {
            'dm_code': dm_code_to_string(dm_code),
            'info_name': representative_section.get('info_name', 'Unknown'),
            'sections': sections_in_group
        }

        print(f"Generated module: {representative_section.get('info_name', 'Unknown')} -> {filename}")

    # Generate Publication Module (PMC)
    pm_filepath = None
    if dm_refs:
        # Use fixed modelIdentCode from first DM or extract short code
        first_dm = dm_refs[0] if dm_refs else None
        model_code = first_dm['dm_code'].get('modelIdentCode', 'S5') if first_dm else 'S5'
        pm_config = create_pm_config(
            model_ident_code=model_code,
            pm_title=document_title,
            enterprise_name=organization,
            originator_name=organization
        )
        pm_filepath = pm_generator.generate_publication_module(pm_config, dm_refs, output_dir, illustrations)
        generated_files.append(pm_filepath)
        print(f"Generated publication module: {os.path.basename(pm_filepath)}")

    # Generate module mapping log
    if split_into_modules:
        from parsers.content_analyzer import generate_module_mapping_log
        mapping_log_path = generate_module_mapping_log(doc_path, module_mapping, output_dir)
    else:
        # For combined mode, generate general module mapping all log
        log_filename = "module_mapping_all.log"
        mapping_log_path = os.path.join(output_dir, log_filename)

        with open(mapping_log_path, 'w', encoding='utf-8') as f:
            import datetime
            f.write("Общее соответствие модулей\n")
            f.write(f"Сгенерировано: {datetime.datetime.now()}\n")
            f.write(f"Исходный документ: {doc_path}\n")
            f.write("=" * 80 + "\n\n")

            # Map source document to generated files
            f.write("Соответствие исходного документа сгенерированным файлам:\n")
            f.write(f"Исходный документ: {os.path.basename(doc_path)}\n\n")

            f.write("Сгенерированные файлы:\n")
            for filename in generated_files:
                basename = os.path.basename(filename)
                if basename.endswith('.xml'):
                    if 'PMC' in basename:
                        f.write(f"  - {basename} (модуль публикации)\n")
                    else:
                        f.write(f"  - {basename} (модуль данных)\n")

            # List graphics files
            graphics_dir = os.path.join(output_dir, "graphics")
            if os.path.exists(graphics_dir):
                graphics_files = [f for f in os.listdir(graphics_dir) if f.endswith('.jpg')]
                for gf in graphics_files:
                    f.write(f"  - {gf} (файл иллюстрации)\n")

            f.write("\n")

            # Detailed module mappings
            for filename, module_data in module_mapping.items():
                dm_code = module_data['dm_code']
                sections = module_data['sections']

                f.write(f"Детали модуля: {filename}\n")
                f.write(f"  Код DM: {dm_code}\n")
                f.write(f"  info_name: {module_data.get('info_name', 'Unknown')}\n")
                f.write("\n")

        print(f"Общий лог соответствия модулей сохранен в: {mapping_log_path}")

    # Validate that all original content is included in generated XML
    validation_errors = validate_content_inclusion(doc_path, generated_files, illustrations, output_dir)

    print(f"\nGenerated {len(generated_files)} files:")
    for filepath in generated_files:
        print(f"  - {os.path.basename(filepath)}")

    if log_path:
        print(f"Content analysis log: {os.path.basename(log_path)}")
    print(f"Elements log: {os.path.basename(elements_log_path)}")
    print(f"Module mapping log: {os.path.basename(mapping_log_path)}")

    if validation_errors:
        print(f"Content validation errors found: {len(validation_errors)}")
        print("See errors.log for details")
    else:
        print("All content successfully validated")

    return generated_files


def group_sections_for_modules(analysis_results: List[Dict], split_into_modules: bool) -> Dict[str, Dict]:
    """
    Group sections that should be combined into the same module.

    Args:
        analysis_results: Results from analyze_document_content
        split_into_modules: Whether to split into multiple modules or combine into one

    Returns:
        Dict mapping group keys to group info with sections
    """
    if not split_into_modules:
        # Combine all sections into a single module
        all_info_names = [section.get('info_name', 'unknown') for section in analysis_results]
        combined_info_name = ', '.join(set(all_info_names)) if all_info_names else 'Combined document'
        return {
            'combined': {
                'sections': analysis_results,
                'info_name': combined_info_name
            }
        }

    # Original logic: split into multiple modules
    groups = {}

    for section in analysis_results:
        section_type = section.get('section_type', 'unknown')
        info_name = section.get('info_name', 'unknown')

        if section_type == 'component':
            # Each component gets its own module
            group_key = f"component_{section.get('start_para', 0)}"
        elif info_name == 'Описание функций изделия':
            # Function descriptions get their own module
            group_key = "function_description"
        elif section_type in ['purpose', 'description', 'operation']:
            # Group sections of same type together (not ideal, but better than duplicates)
            group_key = f"{section_type}_{info_name.replace(' ', '_')}"
        else:
            # Fallback grouping
            group_key = f"misc_{section.get('start_para', 0)}"

        if group_key not in groups:
            groups[group_key] = {
                'sections': [],
                'info_name': info_name
            }

        groups[group_key]['sections'].append(section)

    return groups


def dm_code_to_string(dm_code: Dict) -> str:
    """Convert DM code dict to string representation."""
    return "DMC-S5-A-120-{system_sub}-{sub_sub}-{assy}-{disassy}{disassy_var}-{info}{info_var}-{location}_001".format(
        system_sub=dm_code.get('subSystemCode', '1'),
        sub_sub=dm_code.get('subSubSystemCode', '0'),
        assy=dm_code.get('assyCode', '00'),
        disassy=dm_code.get('disassyCode', '00'),
        disassy_var=dm_code.get('disassyCodeVariant', 'A'),
        info=dm_code.get('infoCode', '012'),
        info_var=dm_code.get('infoCodeVariant', 'A'),
        location=dm_code.get('itemLocationCode', 'A')
    )


def group_sections_by_type(headings: List[str]) -> Dict[str, List[int]]:
    """
    Group section indices by their type for data module creation.

    Args:
        headings: List of section headings

    Returns:
        Dict mapping group types to lists of section indices
    """
    groups = {
        "system_purpose": [],
        "system_description": [],
        "system_operation": [],
        "components": []
    }

    for idx, heading in enumerate(headings):
        heading_lower = heading.lower()

        if 'общие сведения' in heading_lower:
            groups["system_purpose"].append(idx)
        elif 'состав рсуо' in heading_lower or 'оборудовани' in heading_lower or 'структурно представляет' in heading_lower:
            groups["system_description"].append(idx)
        elif 'информацион' in heading_lower or 'режимы работы' in heading_lower or 'управлени' in heading_lower:
            groups["system_operation"].append(idx)
        else:
            # Check if component
            unit_keywords = ['блок', 'пульт', 'рама', 'система', 'контроллер', 'бусп', 'ски', 'бкп', 'пусть', 'пнп', 'би-м', 'сброса', 'псо-5п', 'бус-5пм', 'бч-70']
            if any(keyword in heading_lower for keyword in unit_keywords):
                groups["components"].append(idx)

    return groups


def assemble_content_for_section(section: Dict, document: Document, tables: Dict[str, Dict], lists_data: List[Dict], elements: List[Dict], illustration_counter: int = 0, figure_counter: int = 0, figure_info: List = None, tech_name: str = "", info_name_override: str = "") -> Dict:
    """
    Assemble content for a specific analyzed section using element analysis.

    Args:
        section: Section dictionary from content analysis
        document: Document object
        tables: Dict of table references
        lists_data: List of parsed lists
        elements: List of analyzed elements
        illustration_counter: Current illustration counter
        figure_counter: Current figure counter

    Returns:
        Content dict with xml_parts and updated counters
    """
    xml_parts = []

    # Extract content from the section's paragraph range
    start_para = section.get('start_para', 0)
    end_para = section.get('end_para', len(document.paragraphs) - 1)

    # Get elements in this section
    section_elements = [elem for elem in elements if start_para <= elem.get('start_para', 0) <= end_para]

    # Group consecutive list items
    current_list_items = []
    current_list_type = None

    def flush_current_list():
        nonlocal current_list_items, current_list_type, xml_parts
        if current_list_items:
            # Generate randomList XML wrapped in para to conform to schema
            items_xml = ''.join([f"<listItem><para>{item}</para></listItem>" for item in current_list_items])
            prefix = 'pf02' if current_list_type == 'unnumbered_list' else 'nfp01'
            list_xml = f'<randomList listItemPrefix="{prefix}">{items_xml}</randomList>'
            xml_parts.append(f'<para>{list_xml}</para>')
            current_list_items = []
            current_list_type = None

    # Pre-process elements to clean up table references in paragraphs and handle figure references
    processed_elements = []
    for i, elem in enumerate(section_elements):
        elem_type = elem.get('type', 'paragraph')
        content = elem.get('content', '')

        # Check if this is a paragraph that contains a table reference and is followed by a table
        if elem_type == 'paragraph' and i < len(section_elements) - 1:
            next_elem = section_elements[i + 1]
            if next_elem.get('type') == 'table':
                table_title = next_elem.get('content', '')
                if table_title:
                    # Remove table references from paragraph content
                    # Remove patterns like "(Таблица 1)", "Таблица 1", "ТАБЛИЦА 1", etc.
                    patterns = [
                        r'\s*\([Тт]аблица\s*\d+\)\s*',
                        r'\s*[Тт]аблица\s*\d+\s*',
                        r'\s*[Тт]аб\.\s*\d+\s*',
                        r'\s*[Тт]абл\.\s*\d+\s*'
                    ]
                    for pattern in patterns:
                        content = re.sub(pattern, '', content, flags=re.IGNORECASE)

                    # Clean up extra whitespace and trailing punctuation
                    content = content.strip()
                    if content.endswith('.,'):
                        content = content[:-2]
                    elif content.endswith(','):
                        content = content[:-1]
                    elif content.endswith('.'):
                        pass  # Keep the period
                    else:
                        pass

                    # Update the element content
                    elem = elem.copy()
                    elem['content'] = content

            # Check if this paragraph contains a figure reference and is followed by illustration_reference
            elif next_elem.get('type') == 'illustration_reference':
                # Check if content matches figure name pattern
                figure_name_pattern = r'^[Рр]исунок\s*\d+\s*[–-]\s*.+'
                if re.match(figure_name_pattern, content.strip()):
                    # Mark this element to be skipped (will be used for figure title)
                    elem = elem.copy()
                    elem['skip_output'] = True
                    elem['figure_title'] = content.strip()
                    # Mark the next element with the figure title
                    next_elem_copy = next_elem.copy()
                    next_elem_copy['figure_title'] = content.strip()
                    # Replace the next element in processed_elements
                    processed_elements.append(next_elem_copy)
                    continue  # Skip adding the current elem since it's marked for skip

        processed_elements.append(elem)

    # Track if we're building a levelledPara with multiple elements
    current_levelled_para = []
    in_levelled_para = False
    is_first_paragraph = True

    for elem in processed_elements:
        elem_type = elem.get('type', 'paragraph')
        content = elem.get('content', '')

        # Check if this is the first paragraph and contains both techName and infoName
        if elem_type == 'paragraph' and is_first_paragraph:
            # Use provided tech_name, or fall back to section info
            current_tech_name = tech_name or ""

            content_lower = content.lower()
            tech_name_lower = current_tech_name.lower()

            # Check if techName is in the content
            if tech_name_lower in content_lower:
                # Extract the part after techName as infoName
                tech_name_end = content_lower.find(tech_name_lower) + len(tech_name_lower)
                remaining_content = content_lower[tech_name_end:].strip()

                # Remove common separators like "–", "-", etc.
                separators = [' – ', ' - ', ' –', ' -', '– ', '- ']
                for sep in separators:
                    if remaining_content.startswith(sep):
                        remaining_content = remaining_content[len(sep):].strip()
                        break

                # If we have both techName and extracted infoName, skip the paragraph entirely
                if remaining_content:
                    # Skip outputting this paragraph as it contains both techName and infoName
                    is_first_paragraph = False
                    continue  # Skip adding this element to processed_elements entirely
            is_first_paragraph = False

        if elem_type == 'header':
            # Flush any pending list and levelledPara before adding header
            flush_current_list()
            if current_levelled_para:
                xml_parts.append(f'<levelledPara>{"".join(current_levelled_para)}</levelledPara>')
                current_levelled_para = []
                in_levelled_para = False

            # Generate para element instead of levelledPara
            xml_parts.append(f'<para>{content}</para>')

        elif elem_type == 'numbered_paragraph_header':
            # Always start a new levelledPara for numbered headers
            # Close any pending levelledPara and flush any pending list first
            flush_current_list()
            if current_levelled_para:
                xml_parts.append(f'<levelledPara>{"".join(current_levelled_para)}</levelledPara>')
                current_levelled_para = []
                in_levelled_para = False

            # Strip numbering from the header text for title element
            title_text = re.sub(r'^\d+\.\s*', '', content).strip()

            # Start new levelledPara with title
            current_levelled_para = [f'<title>{title_text}</title>']
            in_levelled_para = True

        elif elem_type in ['numbered_list', 'unnumbered_list']:
            # Flush any pending levelledPara before starting list
            if current_levelled_para:
                xml_parts.append(f'<levelledPara>{"".join(current_levelled_para)}</levelledPara>')
                current_levelled_para = []
                in_levelled_para = False

            # Check if this continues the current list
            if elem_type == current_list_type:
                current_list_items.append(content)
            else:
                # Flush previous list and start new one
                flush_current_list()
                current_list_type = elem_type
                current_list_items = [content]

        elif elem_type == 'paragraph':
            # Add paragraph content to current levelledPara or start new one
            if not in_levelled_para:
                current_levelled_para = []
                in_levelled_para = True

            current_levelled_para.append(f'<para>{content}</para>')

        else:
            # Flush any pending levelledPara and list before adding other elements
            if current_levelled_para:
                xml_parts.append(f'<levelledPara>{"".join(current_levelled_para)}</levelledPara>')
                current_levelled_para = []
                in_levelled_para = False

            flush_current_list()

            if elem_type == 'table':
                # Use the enhanced table XML from element analysis
                table_xml = elem.get('xml_example', '')
                if table_xml:
                    xml_parts.append(table_xml)
            elif elem_type == 'illustration':
                # Generate proper figure with ID and reproduction attributes for embedded illustrations
                figure_id = f"ICN{illustration_counter + 1:02d}"
                graphic_id = f"g{illustration_counter}"
                graphic_ident = f"GS5-A-120-10-00-00A-041A-A_001_RU-RU-GRAPHIC{illustration_counter}"
                # Use the actual content from the element instead of hardcoded text
                figure_title = elem.get('content', 'Название иллюстрации')
                xml_parts.append(f'''<figure id="{figure_id}">
            <title>{figure_title}</title>
            <graphic infoEntityIdent="{graphic_ident}" reproductionScale="32" reproductionWidth="170mm" reproductionHeight="120mm" id="{graphic_id}"/>
          </figure>''')
                if figure_info is not None:
                    figure_info.append({'id': figure_id, 'file': f"GS5-A-120-10-00-00A-041A-A_001_RU-RU-GRAPHIC{illustration_counter}.jpg"})
                illustration_counter += 1
            elif elem_type == 'illustration_reference':
                # For illustration references that are not combined with titles, create internal references
                # Extract reference number from content or details
                ref_number = None
                content = elem.get('content', '')
                details = elem.get('details', '')

                # Try to extract number from content like "Ссылка на иллюстрацию 6"
                match = re.search(r'(\d+)', content)
                if match:
                    ref_number = int(match.group(1))
                else:
                    # Try details
                    match = re.search(r'illustration_reference (\d+)', details)
                    if match:
                        ref_number = int(match.group(1))

                if ref_number:
                    icn_ref = f"ICN{ref_number:02d}"
                    xml_parts.append(f'<para>Ссылка на рисунок: <internalRef internalRefId="{icn_ref}" internalRefTargetType="irtt01"/></para>')
                else:
                    # Fallback
                    xml_parts.append(f'<para>{content}</para>')
            elif elem_type == 'warning':
                xml_parts.append(f'<warning><para>{content}</para></warning>')
            else:
                # Default paragraph - skip if marked for skipping
                if not elem.get('skip_output', False):
                    xml_parts.append(f'<para>{content}</para>')

    # Flush any remaining content
    if current_levelled_para:
        xml_parts.append(f'<levelledPara>{"".join(current_levelled_para)}</levelledPara>')
    flush_current_list()

    return {
        "xml_parts": xml_parts,
        "illustration_counter": illustration_counter,
        "figure_counter": figure_counter
    }


def get_dm_code_for_section(section: Dict, component_counter: int) -> Dict:
    """
    Get DM code for a section based on its type and index.

    Args:
        section: Section dictionary from content analysis
        component_counter: Counter for component numbering

    Returns:
        DM code dictionary
    """
    base_dm_code = {
        'modelIdentCode': 'S5',
        'systemDiffCode': 'A',
        'systemCode': '120',
        'subSystemCode': '1',
        'subSubSystemCode': '0',
        'assyCode': '00',
        'disassyCode': '00',
        'disassyCodeVariant': 'A',
        'infoCode': '011',
        'infoCodeVariant': 'A',
        'itemLocationCode': 'A'
    }

    if section.get('is_component'):
        # Component-specific coding
        base_dm_code.update({
            'subSubSystemCode': f'{component_counter}',
            'infoCode': '017',
            'infoCodeVariant': 'A'
        })

    return base_dm_code


def assemble_content(sections: List[str], text_sections: Dict[str, str], tables: Dict[str, Dict], lists_data: List[Dict]) -> Dict:
    """
    Assemble content from parsed sections.

    Args:
        sections: List of section headings to include
        text_sections: Dict of heading -> text content
        tables: Dict of table references
        lists_data: List of parsed lists

    Returns:
        Content dict with paragraphs, tables, lists
    """
    content = {
        "paragraphs": [],
        "tables": [],
        "lists": []
    }

    # Add paragraphs from sections
    for section in sections:
        if section in text_sections:
            # Split by paragraphs and add
            paras = text_sections[section].split('\n')
            content["paragraphs"].extend([p.strip() for p in paras if p.strip()])

    # Add lists (for now, convert all lists - could be enhanced to map to sections)
    for list_data in lists_data:
        if list_data.get('items'):
            list_xml = convert_list_to_s1000d_randomlist(list_data)
            content["lists"].append(list_xml)

    # Add tables (simplified - add them all)
    from parsers.table_parser import convert_table_to_s1000d_format
    for table_ref, table_data in tables.items():
        table_xml = convert_table_to_s1000d_format(table_data)
        content["tables"].append(table_xml)

    return content


def validate_content_inclusion(doc_path: str, generated_files: List[str], illustrations: Dict[str, str], output_dir: str) -> List[str]:
    """
    Validate that all original document content is included in generated XML files.

    Args:
        doc_path: Path to original document
        generated_files: List of generated XML file paths
        illustrations: Dict of illustrations from original document
        output_dir: Output directory

    Returns:
        List of validation error messages
    """
    errors = []
    error_log_path = os.path.join(output_dir, "errors.log")

    # Load original document
    try:
        doc = Document(doc_path)
    except Exception as e:
        error_msg = f"Failed to load original document {doc_path}: {e}"
        errors.append(error_msg)
        _write_error_log(error_log_path, [error_msg])
        return errors

    # Extract all text content from original document
    original_texts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            original_texts.append(text)

    # Extract all table content
    original_tables = []
    for i, table in enumerate(doc.tables):
        table_content = []
        for row in table.rows:
            row_content = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_content.append(cell_text)
            if row_content:
                table_content.append(row_content)
        if table_content:
            original_tables.append({
                'index': i,
                'content': table_content
            })

    # Extract illustration references from document text
    original_illustration_refs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        # Find figure references
        figure_matches = re.findall(r'[Рр]исунок\s*\d+', text, re.IGNORECASE)
        original_illustration_refs.extend(figure_matches)

    # Remove duplicates
    original_illustration_refs = list(set(original_illustration_refs))

    # Check generated XML files for content inclusion
    found_texts = set()
    found_tables = set()
    found_illustrations = set()

    for xml_file in generated_files:
        if not xml_file.endswith('.xml'):
            continue

        try:
            with open(xml_file, 'r', encoding='utf-8') as f:
                xml_content = f.read()

            # Extract text content from XML (from <para> tags)
            para_texts = re.findall(r'<para[^>]*>(.*?)</para>', xml_content, re.DOTALL)
            for para_text in para_texts:
                # Clean XML tags from text
                clean_text = re.sub(r'<[^>]+>', '', para_text).strip()
                if clean_text:
                    found_texts.add(clean_text)

            # Extract table content from XML
            table_matches = re.findall(r'<table[^>]*>.*?</table>', xml_content, re.DOTALL)
            for table_xml in table_matches:
                found_tables.add(table_xml.strip())

            # Extract illustration references from XML
            figure_titles = re.findall(r'<title>(.*?)</title>', xml_content)
            for title in figure_titles:
                if 'рисунок' in title.lower() or 'иллюстрация' in title.lower():
                    found_illustrations.add(title.strip())

        except Exception as e:
            error_msg = f"Failed to read generated XML file {xml_file}: {e}"
            errors.append(error_msg)
            continue

    # Validate text content inclusion
    missing_texts = []
    for original_text in original_texts:
        # Check if this text (or significant portion) is found in generated XML
        text_found = False
        for found_text in found_texts:
            # Use fuzzy matching - check if original text is contained in found text
            # or if found text is contained in original text (for partial matches)
            if (original_text in found_text or
                found_text in original_text or
                _similar_text(original_text, found_text, threshold=0.8)):
                text_found = True
                break

        if not text_found and len(original_text) > 10:  # Only report missing texts longer than 10 chars
            missing_texts.append(original_text)

    # Validate table inclusion
    missing_tables = []
    for table_info in original_tables:
        table_found = False
        table_content_str = str(table_info['content'])

        for found_table in found_tables:
            if table_content_str in found_table:
                table_found = True
                break

        if not table_found:
            missing_tables.append(f"Table {table_info['index'] + 1}: {table_content_str[:200]}...")

    # Validate illustration inclusion
    missing_illustrations = []
    for ill_ref in original_illustration_refs:
        ill_found = False
        for found_ill in found_illustrations:
            if ill_ref.lower() in found_ill.lower():
                ill_found = True
                break

        if not ill_found:
            missing_illustrations.append(ill_ref)

    # Validate illustrations dict
    missing_illustration_files = []
    graphics_dir = os.path.join(output_dir, "graphics")
    if os.path.exists(graphics_dir):
        existing_graphics = set(f for f in os.listdir(graphics_dir) if f.endswith('.jpg'))
        for ill_key, ill_file in illustrations.items():
            if ill_file not in existing_graphics:
                missing_illustration_files.append(ill_file)

    # Compile all errors
    if missing_texts:
        errors.append(f"Missing text content ({len(missing_texts)} items)")
        for i, text in enumerate(missing_texts[:10]):  # Limit to first 10
            errors.append(f"  Missing text {i+1}: {text[:100]}{'...' if len(text) > 100 else ''}")
        if len(missing_texts) > 10:
            errors.append(f"  ... and {len(missing_texts) - 10} more missing texts")

    if missing_tables:
        errors.append(f"Missing tables ({len(missing_tables)} items)")
        for table in missing_tables[:5]:  # Limit to first 5
            errors.append(f"  {table}")

    if missing_illustrations:
        errors.append(f"Missing illustration references ({len(missing_illustrations)} items)")
        for ill in missing_illustrations:
            errors.append(f"  {ill}")

    if missing_illustration_files:
        errors.append(f"Missing illustration files ({len(missing_illustration_files)} items)")
        for ill_file in missing_illustration_files:
            errors.append(f"  {ill_file}")

    # Write detailed error log
    _write_error_log(error_log_path, errors)

    return errors


def _similar_text(text1: str, text2: str, threshold: float = 0.8) -> bool:
    """
    Check if two texts are similar based on word overlap.

    Args:
        text1: First text
        text2: Second text
        threshold: Similarity threshold (0.0 to 1.0)

    Returns:
        True if texts are similar
    """
    words1 = set(text1.lower().split())
    words2 = set(text2.lower().split())

    if not words1 or not words2:
        return False

    intersection = words1.intersection(words2)
    union = words1.union(words2)

    similarity = len(intersection) / len(union)
    return similarity >= threshold


def _text_similar(text1: str, text2: str, threshold: float = 0.6) -> bool:
    """
    Check if two texts are similar, allowing for partial matches.

    Args:
        text1: First text
        text2: Second text
        threshold: Similarity threshold (0.0 to 1.0)

    Returns:
        True if texts are similar
    """
    if not text1 or not text2:
        return False

    text1_lower = text1.lower()
    text2_lower = text2.lower()

    # Check if one text contains the other
    if text1_lower in text2_lower or text2_lower in text1_lower:
        return True

    # Check word overlap
    words1 = set(text1_lower.split())
    words2 = set(text2_lower.split())

    if not words1 or not words2:
        return False

    intersection = words1.intersection(words2)
    union = words1.union(words2)

    similarity = len(intersection) / len(union)
    return similarity >= threshold


def _write_error_log(log_path: str, errors: List[str]):
    """
    Write validation errors to log file.

    Args:
        log_path: Path to error log file
        errors: List of error messages
    """
    try:
        with open(log_path, 'w', encoding='utf-8') as f:
            f.write("Content Validation Errors\n")
            f.write("=" * 50 + "\n")
            f.write(f"Generated: {os.path.basename(__file__)} at {os.path.getctime(__file__)}\n\n")

            if not errors:
                f.write("No validation errors found. All content appears to be included.\n")
            else:
                f.write(f"Found {len(errors)} validation errors:\n\n")
                for i, error in enumerate(errors, 1):
                    f.write(f"{i}. {error}\n")

                f.write("\n" + "=" * 50 + "\n")
                f.write("Detailed Analysis:\n")
                f.write("- Text content: Checked all paragraphs from original document\n")
                f.write("- Tables: Checked table structures and content\n")
                f.write("- Illustrations: Checked figure titles and references\n")
                f.write("- Files: Checked existence of illustration files\n")

    except Exception as e:
        print(f"Failed to write error log: {e}")
