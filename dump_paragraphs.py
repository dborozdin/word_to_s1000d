#!/usr/bin/env python3
"""
Script to dump paragraphs from the test document to see actual content
"""

import os
from docx import Document

def dump_paragraphs(doc_path: str, output_file: str, max_paragraphs: int = 50):
    """
    Dump first N paragraphs from document to see actual content
    """
    if not os.path.exists(doc_path):
        print(f"Document not found: {doc_path}")
        return

    doc = Document(doc_path)
    print(f"Document loaded. Total paragraphs: {len(doc.paragraphs)}")

    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(f"First {max_paragraphs} paragraphs from {os.path.basename(doc_path)}\n")
        f.write("=" * 80 + "\n\n")

        for i, paragraph in enumerate(doc.paragraphs[:max_paragraphs]):
            text = paragraph.text.strip()
            style = paragraph.style.name if paragraph.style else 'Normal'

            f.write(f"Paragraph {i}:\n")
            f.write(f"  Style: {style}\n")
            f.write(f"  Text: '{text}'\n")

            # Check if it starts with number
            if text and text[0].isdigit():
                f.write(f"  STARTS WITH DIGIT: Yes\n")

            # Check for number patterns
            import re
            if re.match(r'^\s*\d+\.\s*', text):
                f.write(f"  NUMBER PATTERN: Yes\n")

            f.write("\n")

    print(f"Dump saved to {output_file}")

if __name__ == "__main__":
    dump_paragraphs("docs/РСУ_адаптированная.docx", "paragraphs_dump.txt", 100)
