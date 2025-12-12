"""
Table parser for docx documents.
Extracts table data for conversion to S1000D format.
"""

from typing import Dict, List, Tuple
from docx import Document


def extract_tables(doc: Document) -> List[Dict[str, List[List[str]]]]:
    """
    Extract all tables from document with metadata.

    Args:
        doc: Docx document object

    Returns:
        List of tables, each as dict with 'rows' key containing list of lists
    """
    tables_data = []

    for table in doc.tables:
        table_data = {'rows': []}

        # Extract rows
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data['rows'].append(row_data)

        tables_data.append(table_data)

    return tables_data


def get_tables_by_reference(doc: Document, table_references: Dict[int, str] = None) -> Dict[str, Dict[str, List[List[str]]]]:
    """
    Extract tables mapped to their reference names.

    Args:
        doc: Docx document object
        table_references: Optional mapping of table index to reference name

    Returns:
        Dictionary of table reference -> table data
    """
    if table_references is None:
        # Use "Table 1", "Table 2", etc.
        table_references = {i: f"Table {i+1}" for i in range(len(doc.tables))}

    tables_dict = {}

    for idx, table in enumerate(doc.tables):
        ref_name = table_references.get(idx, f"Table {idx+1}")

        table_data = {'rows': []}
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data['rows'].append(row_data)

        tables_dict[ref_name] = table_data

    return tables_dict


def extract_enhanced_tables_with_titles(doc: Document) -> List[Dict]:
    """
    Extract tables with titles from paragraphs before each table.

    Args:
        doc: Docx document object

    Returns:
        List of enhanced table data with titles, headers, and content
    """
    tables_data = []

    # Get all document elements (paragraphs and tables) in order
    doc_elements = []
    para_idx = 0
    table_idx = 0

    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            if para_idx < len(doc.paragraphs):
                doc_elements.append(('paragraph', para_idx))
                para_idx += 1
        elif element.tag.endswith('tbl'):  # Table
            if table_idx < len(doc.tables):
                doc_elements.append(('table', table_idx))
                table_idx += 1

    # Process elements to find table titles
    for i, (element_type, element_idx) in enumerate(doc_elements):
        if element_type == 'table':
            table = doc.tables[element_idx]

            # Look for title in previous paragraph
            title = ""
            if i > 0 and doc_elements[i-1][0] == 'paragraph':
                prev_para = doc.paragraphs[doc_elements[i-1][1]]
                prev_text = prev_para.text.strip()
                # Look for "Таблица" or "Таб" followed by number
                import re
                match = re.search(r'(?:[Тт]аблица|[Тт]аб)\s*\d+', prev_text, re.IGNORECASE)
                if match:
                    title = match.group(0)
                else:
                    # Try to find any text that looks like a table reference
                    match = re.search(r'.*[Тт]аблица.*', prev_text, re.IGNORECASE)
                    if match:
                        title = match.group(0).strip()

            # Extract table rows
            rows = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                rows.append(row_data)

            # Create enhanced table data
            table_data = {
                'id': f'tab-{element_idx + 1}',
                'title': title,
                'rows': rows,
                'headers': rows[0] if rows else [],  # First row as headers
                'body_rows': rows[1:] if len(rows) > 1 else []  # Rest as body
            }

            tables_data.append(table_data)

    return tables_data


def convert_enhanced_table_to_s1000d_format(table_data: Dict) -> str:
    """
    Convert enhanced table data to S1000D table XML structure with ID, title, and headers.

    Args:
        table_data: Enhanced table data with id, title, headers, body_rows

    Returns:
        Enhanced XML string representation of table
    """
    table_id = table_data.get('id', '')
    title = table_data.get('title', '')
    headers = table_data.get('headers', [])
    body_rows = table_data.get('body_rows', [])

    if not headers and not body_rows:
        return ""

    num_cols = max(len(headers), max(len(row) for row in body_rows) if body_rows else 0)

    # Build colspec elements
    colspecs = []
    for i in range(num_cols):
        colspecs.append(f'<colspec colnum="{i+1}" colname="col{i+1}" colwidth="1*"/>')

    # Build thead section
    thead_xml = ""
    if headers:
        header_entries = []
        for col_idx, header_text in enumerate(headers):
            if header_text.strip():
                header_entries.append(f'<entry colname="col{col_idx+1}"><para>{header_text}</para></entry>')
            else:
                header_entries.append(f'<entry colname="col{col_idx+1}"></entry>')

        # Pad missing columns if needed
        while len(header_entries) < num_cols:
            header_entries.append(f'<entry colname="col{len(header_entries)+1}"></entry>')

        thead_xml = f'<thead><row rowsep="1">{''.join(header_entries)}</row></thead>'

    # Build tbody section
    tbody_entries = []
    all_rows = body_rows if body_rows else [[]]

    for row in all_rows:
        row_entries = []
        for col_idx, cell_text in enumerate(row):
            if cell_text.strip():
                row_entries.append(f'<entry colname="col{col_idx+1}"><para>{cell_text}</para></entry>')
            else:
                row_entries.append(f'<entry colname="col{col_idx+1}"></entry>')

        # Pad missing columns if needed
        while len(row_entries) < num_cols:
            row_entries.append(f'<entry colname="col{len(row_entries)+1}"></entry>')

        tbody_entries.append(f'<row>{''.join(row_entries)}</row>')

    # Build complete table XML
    title_xml = f'<title>{title}</title>' if title else ''
    table_xml = f'''<table id="{table_id}" frame="topbot" colsep="0" rowsep="0">
{title_xml}
<tgroup cols="{num_cols}">
{chr(10).join(colspecs)}
{thead_xml}
<tbody>
{chr(10).join(tbody_entries)}
</tbody>
</tgroup>
</table>'''

    return table_xml


def convert_table_to_s1000d_format(table_data: Dict[str, List[List[str]]]) -> str:
    """
    Convert table data to basic S1000D table XML structure.

    Args:
        table_data: Table data as rows of lists

    Returns:
        Basic XML string representation of table
    """
    if not table_data.get('rows'):
        return ""

    rows = table_data['rows']
    num_cols = max(len(row) for row in rows) if rows else 0

    # Build basic table entry
    table_entries = []

    for row in rows:
        row_entries = []
        for col_idx, cell in enumerate(row):
            if cell.strip():
                row_entries.append(f"<entry><para>{cell}</para></entry>")
            else:
                row_entries.append("<entry></entry>")
        # Pad missing columns if needed
        while len(row_entries) < num_cols:
            row_entries.append("<entry></entry>")

        table_entries.append(f"<row>{''.join(row_entries)}</row>")

    return f"<table><tgroup cols=\"{num_cols}\"><tbody>{''.join(table_entries)}</tbody></tgroup></table>"
