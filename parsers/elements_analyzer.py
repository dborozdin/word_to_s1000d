"""
Element analyzer for Word documents.
Analyzes document structure and generates logs of all detected elements.

Parsing heuristics are documented in parsing_rules.json at project root.
"""

import os
import re
import json
import hashlib
from typing import Dict, List, Optional, Tuple, Any
from docx import Document
from docx.shared import Inches
import datetime


def compute_stable_id(seq_index: int, element_type: str, text: str) -> str:
    """Generate a stable 12-hex element identifier from DOCX parsing.

    Uses sequential index + type + text prefix to ensure uniqueness
    even for duplicate/similar paragraphs.
    """
    prefix = text[:80].strip().lower() if text else ""
    raw = f"{seq_index}|{element_type}|{prefix}"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:12]


def get_parsing_rules() -> dict:
    """Load parsing rules from parsing_rules.json."""
    from app_paths import get_parsing_rules_path
    rules_path = get_parsing_rules_path()
    if os.path.isfile(rules_path):
        with open(rules_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}


def analyze_document_elements(doc: Document, illustrations: Dict[str, str] = None, illustration_positions: Dict[str, Dict] = None, llm_config: Dict[str, Any] = None, graphic_ident_prefix: str = None) -> List[Dict[str, Any]]:
    """
    Analyze document and extract all elements with their start/end positions.

    Args:
        doc: Word document object
        illustrations: Dictionary of illustration references
        illustration_positions: Dictionary of illustration positions
        llm_config: Optional LLM configuration for element classification

    Returns:
        List of element dictionaries with type, start_para, end_para, content, xml_example
    """
    # Initialize LLM classifier if enabled
    llm_classifier = None
    if llm_config and llm_config.get('enabled', False):
        try:
            from llm import DocumentStructureClassifier
            llm_classifier = DocumentStructureClassifier(llm_config)
            if llm_classifier.is_llm_available():
                print("[LLM] Structure classifier initialized")
            else:
                print("[LLM] Ollama not available, using heuristics only")
                llm_classifier = None
        except ImportError as e:
            print(f"[LLM] Failed to import classifier: {e}")
            llm_classifier = None
    elements = []

    # Graphic naming prefix (dynamic per DM or hardcoded fallback)
    effective_prefix = graphic_ident_prefix or "GS5-A-120-10-00-00A-041A-A_001_RU-RU"

    # Track list state
    current_list = None
    list_start_para = 0
    list_base_level = 0  # Base indentation of current list (for nested detection)

    # Track numbering state for numbered paragraph headers
    # This tracks the actual numbering sequence as it appears in the document
    numbering_counters = {}  # numId -> {ilvl -> counter}
    global_numbering_counter = 1  # Simple sequential counter for basic cases

    # Reset ПУНКТ-style counters for this document
    global _punkt_counters
    _punkt_counters = {}

    # Track cumulative text for position calculation
    line_number = 1
    char_position = 0

    # Process both paragraphs and tables in document order
    # Get all document elements (paragraphs and tables) in order
    doc_elements = []
    para_idx = 0
    table_idx = 0

    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            if para_idx < len(doc.paragraphs):
                doc_elements.append(('paragraph', para_idx))
                para_idx += 1
        elif element.tag.endswith('tbl'):  # Table
            if table_idx < len(doc.tables):
                doc_elements.append(('table', table_idx))
                table_idx += 1

    # Extract enhanced table data with titles
    from parsers.table_parser import extract_enhanced_tables_with_titles
    enhanced_tables = extract_enhanced_tables_with_titles(doc)

    # Process elements in document order
    current_table_idx = 0
    illustration_counter = 0  # Track illustration numbering for infoEntityIdent
    for i, (element_type, element_idx) in enumerate(doc_elements):
        if element_type == 'table':
            # Handle table element with enhanced data
            table_data = enhanced_tables[current_table_idx] if current_table_idx < len(enhanced_tables) else None
            current_table_idx += 1

            # Update position tracking for table
            table_start_line = line_number
            table_start_char = char_position

            # Approximate table size (simplified - assume 1 line per row)
            table_rows = len(table_data['rows']) if table_data else 1
            line_number += table_rows
            char_position = 0  # Reset to start of line

            table_end_line = line_number
            table_end_char = char_position

            # Create enhanced table XML
            if table_data:
                from parsers.table_parser import convert_enhanced_table_to_s1000d_format
                table_xml = convert_enhanced_table_to_s1000d_format(table_data)
            else:
                table_xml = '<table><tgroup cols="2"><tbody><row><entry>Ячейка 1</entry><entry>Ячейка 2</entry></row></tbody></tgroup></table>'

            # Create table element
            element_info = {
                'type': 'table',
                'start_line': table_start_line,
                'start_char': table_start_char,
                'end_line': table_end_line,
                'end_char': table_end_char,
                'start_para': para_idx,  # Use current para_idx as approximation
                'end_para': para_idx,
                'content': table_data.get('title', f'Таблица {element_idx + 1}') if table_data else f'Таблица {element_idx + 1}',
                'xml_example': table_xml,
                'details': f'Таблица {element_idx + 1}'
            }
            elements.append(element_info)
            continue
        else:
            # Handle paragraph element
            paragraph = doc.paragraphs[element_idx]
            para_idx = element_idx
            text = paragraph.text.strip()
            style_name = paragraph.style.name

            # Check for numbering properties
            numPr = None
            if hasattr(paragraph.paragraph_format, 'numPr') and paragraph.paragraph_format.numPr is not None:
                numPr = paragraph.paragraph_format.numPr
            else:
                # Fallback: check raw OOXML for numPr when python-docx fails to detect it
                # This happens with some DOCX files where numPr is present but not detected by python-docx
                try:
                    pPr_elem = paragraph._element.find('.//w:pPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    if pPr_elem is not None:
                        numPr_elem = pPr_elem.find('.//w:numPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                        if numPr_elem is not None:
                            # Create a simple object to mimic numPr
                            class MockNumPr:
                                def __init__(self, numId_val, ilvl_val):
                                    self.numId_val = numId_val
                                    self.ilvl_val = ilvl_val

                                @property
                                def numId(self):
                                    class MockNumId:
                                        def __init__(self, val):
                                            self.val = val
                                    return MockNumId(self.numId_val)

                                @property
                                def ilvl(self):
                                    class MockIlvl:
                                        def __init__(self, val):
                                            self.val = val
                                    return MockIlvl(self.ilvl_val)

                            # Extract values from raw XML
                            ilvl_elem = numPr_elem.find('.//w:ilvl', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                            numId_elem = numPr_elem.find('.//w:numId', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

                            ilvl_val = int(ilvl_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')) if ilvl_elem is not None else 0
                            numId_val = int(numId_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')) if numId_elem is not None else 0

                            if numId_val > 0:  # Only consider it valid if numId is present
                                numPr = MockNumPr(numId_val, ilvl_val)
                except Exception as e:
                    # If anything fails, just continue without numPr
                    # print(f"FALLBACK FAILED for para {para_idx}: {e}")
                    pass

            # Update position tracking
            para_start_line = line_number
            para_start_char = char_position

            # Add paragraph text to cumulative text (including original formatting for position tracking)
            original_text = paragraph.text

            # Count newlines and characters
            newline_count = original_text.count('\n')
            line_number += newline_count
            if newline_count > 0:
                char_position = len(original_text.split('\n')[-1])
            else:
                char_position += len(original_text)

            para_end_line = line_number
            para_end_char = char_position

            if not text:
                continue

        # PHASE 1 FIX: Check for list markers BEFORE header detection
        # This prevents list items from being incorrectly classified as headers
        is_likely_list_item = _is_likely_list_item(text, doc.paragraphs, para_idx, elements)

        # PHASE 2: Use LLM classifier for ambiguous cases
        llm_says_header = False  # Currently unused, reserved for future enhancements
        llm_says_paragraph = False
        if llm_classifier and not is_likely_list_item:
            # Only use LLM for ambiguous cases (medium length text without clear markers)
            if 20 < len(text) < 300 and not text.endswith(';'):
                context = {
                    'prev_element': elements[-1] if elements else {},
                    'prev_text_ending': elements[-1].get('content', '')[-1:] if elements else '',
                    'style_name': style_name
                }
                classification = llm_classifier.classify_element(text, context)
                if classification and classification.confidence >= llm_classifier.confidence_threshold:
                    if classification.element_type == 'list_item':
                        is_likely_list_item = True
                    elif classification.element_type == 'header':
                        llm_says_header = True
                    elif classification.element_type == 'paragraph':
                        llm_says_paragraph = True

        # Detect numbered paragraph headers first
        # (paragraphs with numbering properties that appear as numbered headings)
        is_numbered_header = False
        level = 1

        # Skip header detection if this looks like a list item or LLM says it's a paragraph
        if numPr is not None and not is_likely_list_item and not llm_says_paragraph:
            # Check if this has numbering and looks like a header
            # Either has Heading style, or has specific header-like content, or just has numbering
            if (style_name.startswith('Heading') or
                style_name in ['1', '2', '3', '4', '5', '6', '7', '8', '9'] or  # Numeric style names like "2"
                # More inclusive check: short text without sentence-ending punctuation
                (len(text.strip()) < 200 and not text.endswith(('.', '!', '?')) and
                 not any(word in text.lower() for word in ['представляет', 'обеспечивает', 'осуществляет', 'является', 'который', 'которая', 'которое', 'которые'])) or
                # Or if it starts with typical header words
                any(text.strip().startswith(word) for word in ['ОБЩИЕ', 'СОСТАВ', 'ОПИСАНИЕ', 'Пульт', 'Блок', 'Рама', 'Таблица', 'Рисунок'])):
                is_numbered_header = True

                # Extract level from style name if possible
                if style_name.startswith('Heading'):
                    level_match = re.search(r'Heading\s*(\d+)', style_name, re.IGNORECASE)
                    level = int(level_match.group(1)) if level_match else 1
                elif style_name.isdigit():
                    level = int(style_name)

        if is_numbered_header:
            # Extract the header part and any remaining content
            header_text = ""
            remaining_content = ""

            # Try to split on newlines first
            if '\n' in text:
                lines = text.split('\n', 1)
                header_text = lines[0].strip()
                remaining_content = lines[1].strip() if len(lines) > 1 else ""
            else:
                # No newlines, try to identify header patterns
                # Look for common header patterns followed by content
                header_patterns = [
                    (r'^([ОБЩИЕ СВЕДЕНИЯ]+)\s+(.+)', 'ОБЩИЕ СВЕДЕНИЯ'),
                    (r'^([СОСТАВ РСУО]+)\s+(.+)', 'СОСТАВ РСУО'),
                    (r'^([АВТОМАТИЧЕСКИЙ РЕЖИМ РАБОТЫ РСУО]+)\s+(.+)', 'АВТОМАТИЧЕСКИЙ РЕЖИМ РАБОТЫ РСУО'),
                    (r'^([АВТОНОМНЫЙ РЕЖИМ РАБОТЫ РСУО]+)\s+(.+)', 'АВТОНОМНЫЙ РЕЖИМ РАБОТЫ РСУО'),
                    (r'^([АВАРИЙНЫЙ РЕЖИМ РАБОТЫ РСУО]+)\s+(.+)', 'АВАРИЙНЫЙ РЕЖИМ РАБОТЫ РСУО'),
                    (r'^([УЧЕБНО-ТРЕНИРОВОЧНЫЙ РЕЖИМ РАБОТЫ РСУО]+)\s+(.+)', 'УЧЕБНО-ТРЕНИРОВОЧНЫЙ РЕЖИМ РАБОТЫ РСУО'),
                    (r'^([РЕЖИМ УПРАВЛЕНИЯ СТВОРКАМИ ГРУЗОВЫХ ОТСЕКОВ]+)\s+(.+)', 'РЕЖИМ УПРАВЛЕНИЯ СТВОРКАМИ ГРУЗОВЫХ ОТСЕКОВ'),
                ]

                for pattern, expected_header in header_patterns:
                    match = re.match(pattern, text, re.IGNORECASE)
                    if match and match.group(1).strip().upper() == expected_header.upper():
                        header_text = match.group(1).strip()
                        remaining_content = match.group(2).strip()
                        break

                # If no pattern matched, assume the whole text is header (fallback)
                if not header_text:
                    header_text = text.strip()

            # Get numbering values
            ilvl_val = 0
            numId_val = 0
            if hasattr(numPr, 'ilvl') and hasattr(numPr, 'numId'):
                ilvl_val = numPr.ilvl.val if hasattr(numPr.ilvl, 'val') else numPr.ilvl
                numId_val = numPr.numId.val if hasattr(numPr.numId, 'val') else numPr.numId

            # Track numbering state and get the correct number
            if numId_val not in numbering_counters:
                numbering_counters[numId_val] = {}
            if ilvl_val not in numbering_counters[numId_val]:
                numbering_counters[numId_val][ilvl_val] = 0

            # Increment counter for this level
            numbering_counters[numId_val][ilvl_val] += 1
            current_number = numbering_counters[numId_val][ilvl_val]

            # Create numbering prefix based on level
            numbering_prefix = f"{current_number}."

            # Combine numbering with header text
            numbered_title = f"{numbering_prefix} {header_text}".strip()

            # Add the numbered paragraph header
            element_info = {
                'type': 'numbered_paragraph_header',
                'start_line': para_start_line,
                'start_char': para_start_char,
                'end_line': para_end_line if not remaining_content else para_start_line,  # Adjust end line if content follows
                'end_char': para_end_char if not remaining_content else para_start_char + len(header_text),
                'start_para': para_idx,
                'end_para': para_idx,
                'content': numbered_title,  # Include numbering in content
                'xml_example': f'<levelledPara><title>{numbered_title}</title></levelledPara>',
                'details': f'Нумерованный заголовок параграфа (уровень {level})'
            }
            elements.append(element_info)

            # If there's remaining content, add it as a separate paragraph
            if remaining_content:
                para_content_start_line = para_start_line  # Same line
                para_content_start_char = para_start_char + len(header_text) + 1  # After header + space

                element_info_content = {
                    'type': 'paragraph',
                    'start_line': para_content_start_line,
                    'start_char': para_content_start_char,
                    'end_line': para_end_line,
                    'end_char': para_end_char,
                    'start_para': para_idx,
                    'end_para': para_idx,
                    'content': remaining_content,
                    'xml_example': f'<para>{remaining_content}</para>',
                    'details': 'Параграф'
                }
                elements.append(element_info_content)

            continue

        # Detect regular headers/headings - only treat as header if it looks like a proper section header
        # PHASE 1 FIX: Skip if this looks like a list item
        # PHASE 2: Also skip if LLM says it's a paragraph
        if style_name.startswith('Heading') and not is_likely_list_item and not llm_says_paragraph:
            level_match = re.search(r'Heading\s*(\d+)', style_name, re.IGNORECASE)
            level = int(level_match.group(1)) if level_match else 1

            # Check if this looks like a real section header (short, title-like text)
            # If it's long or contains sentence structure, treat as regular paragraph
            is_real_header = (
                len(text) < 100 and  # Short text
                not text.endswith(':') and  # Not ending with colon (descriptive)
                not any(word in text.lower() for word in ['представляет', 'обеспечивает', 'осуществляет', 'является', 'таблица', 'рисунок', 'указан'])  # Not descriptive sentences or references
            )

            if not is_real_header:
                # Treat as regular paragraph instead
                pass  # Fall through to paragraph processing
            elif '\n' in text:
                lines = text.split('\n', 1)
                title = lines[0].strip()
                body = lines[1].strip()
                # Add header
                element_info = {
                    'type': 'header',
                    'start_line': para_start_line,
                    'start_char': para_start_char,
                    'end_line': para_start_line,  # Header ends at first line
                    'end_char': para_start_char + len(lines[0]),
                    'start_para': para_idx,
                    'end_para': para_idx,
                    'content': title,
                    'xml_example': f'<levelledPara><title>{title}</title></levelledPara>',
                    'details': f'Уровень {level}'
                }
                elements.append(element_info)
                # Add paragraph for the rest
                if body:
                    para_start_line_body = para_start_line + 1
                    para_start_char_body = 0  # Assume new line
                    element_info_para = {
                        'type': 'paragraph',
                        'start_line': para_start_line_body,
                        'start_char': para_start_char_body,
                        'end_line': para_end_line,
                        'end_char': para_end_char,
                        'start_para': para_idx,
                        'end_para': para_idx,
                        'content': body,
                        'xml_example': f'<para>{body}</para>',
                        'details': 'Параграф'
                    }
                    elements.append(element_info_para)
                continue
            else:
                element_info = {
                    'type': 'header',
                    'start_line': para_start_line,
                    'start_char': para_start_char,
                    'end_line': para_end_line,
                    'end_char': para_end_char,
                    'start_para': para_idx,
                    'end_para': para_idx,
                    'content': text,
                    'xml_example': f'<levelledPara><title>{text}</title></levelledPara>',
                    'details': f'Уровень {level}'
                }
                elements.append(element_info)
                continue

        # Detect main section headers (like 1., 2., etc.) - treat as numbered paragraph headers
        # PHASE 1 FIX: Skip if this looks like a list item
        # PHASE 2: Also skip if LLM says it's a paragraph
        if _is_main_section_header(text) and not is_likely_list_item and not llm_says_paragraph:
            element_info = {
                'type': 'numbered_paragraph_header',
                'start_line': para_start_line,
                'start_char': para_start_char,
                'end_line': para_end_line,
                'end_char': para_end_char,
                'start_para': para_idx,
                'end_para': para_idx,
                'content': text,
                'xml_example': f'<levelledPara><title>{text}</title></levelledPara>',
                'details': 'Нумерованный заголовок параграфа'
            }
            elements.append(element_info)
            continue

        # Detect ПУНКТ-style headers/items (custom Word styles with auto-numbering)
        # Must be before list detection — ПУНКТ paragraphs may have numPr
        # which would cause them to be misclassified as list items.
        # Levels 1-2 → numbered_paragraph_header, level 3+ → numbered_list
        punkt_info = _parse_punkt_style(style_name)
        if punkt_info:
            punkt_level = punkt_info[0]
            numbering = _get_punkt_number(punkt_level)
            numbered_content = f"{numbering} {text}"

            if punkt_level <= 2:
                # Levels 1-2: section headers (e.g. "1 Меры безопасности", "3.2 Последовательность монтажа")
                element_info = {
                    'type': 'numbered_paragraph_header',
                    'start_line': para_start_line,
                    'start_char': para_start_char,
                    'end_line': para_end_line,
                    'end_char': para_end_char,
                    'start_para': para_idx,
                    'end_para': para_idx,
                    'content': numbered_content,
                    'xml_example': f'<levelledPara><title>{numbered_content}</title></levelledPara>',
                    'details': f'Нумерованный заголовок параграфа (уровень {punkt_level}, стиль ПУНКТ)',
                    'numbering_source': 'punkt_style'
                }
            else:
                # Level 3+: numbered list items (e.g. "3.2.1 Снимите транспортные заглушки")
                xml_example = f'<randomList listItemPrefix="pf01"><listItem><para>{numbered_content}</para></listItem></randomList>'
                element_info = {
                    'type': 'numbered_list',
                    'start_line': para_start_line,
                    'start_char': para_start_char,
                    'end_line': para_end_line,
                    'end_char': para_end_char,
                    'start_para': para_idx,
                    'end_para': para_idx,
                    'content': numbered_content,
                    'xml_example': xml_example,
                    'details': f'Нумерованный пункт (уровень {punkt_level}, стиль ПУНКТ)',
                    'numbering_source': 'punkt_style'
                }
            elements.append(element_info)
            continue

        # Style-based warning detection (e.g. "07. Внимание для п.1")
        if _is_warning_style(style_name):
            element_info = {
                'type': 'warning',
                'start_line': para_start_line,
                'start_char': para_start_char,
                'end_line': para_end_line,
                'end_char': para_end_char,
                'start_para': para_idx,
                'end_para': para_idx,
                'content': text,
                'xml_example': '<warning><warningAndCautionPara>' + text + '</warningAndCautionPara></warning>',
                'details': f'Предупреждение (стиль: {style_name})'
            }
            elements.append(element_info)
            continue

        # Style-based note detection (e.g. "07. Замеч для п.1")
        if _is_note_style(style_name):
            element_info = {
                'type': 'note',
                'start_line': para_start_line,
                'start_char': para_start_char,
                'end_line': para_end_line,
                'end_char': para_end_char,
                'start_para': para_idx,
                'end_para': para_idx,
                'content': text,
                'xml_example': '<note><notePara>' + text + '</notePara></note>',
                'details': f'Примечание (стиль: {style_name})'
            }
            elements.append(element_info)
            continue

        # Detect lists
        list_type = _get_list_type(paragraph, para_idx)
        if list_type:
            item_level = _get_list_level(paragraph)

            if current_list != list_type:
                # End previous list if any
                if current_list:
                    # Update end_para for previous list items
                    _update_list_end_para(elements, para_idx - 1)

                # Start new list
                current_list = list_type
                list_start_para = para_idx
                list_base_level = item_level  # Track base indentation of this list

            # Auto-detect nested list: if item level is greater than the
            # base level of the current list, mark as nested type
            effective_type = list_type
            if item_level > list_base_level:
                effective_type = 'nested_' + list_type

            # Add list item
            clean_text = _clean_list_item_text(paragraph, list_type)
            xml_prefix = 'pf02' if list_type == 'unnumbered_list' else 'pf01'
            xml_example = f'<randomList listItemPrefix="{xml_prefix}"><listItem><para>{clean_text}</para></listItem></randomList>'

            element_info = {
                'type': effective_type,
                'start_line': para_start_line,
                'start_char': para_start_char,
                'end_line': para_end_line,
                'end_char': para_end_char,
                'start_para': para_idx,
                'end_para': para_idx,  # Will be updated when list ends
                'content': clean_text,
                'xml_example': xml_example,
                'details': f'Элемент списка ({effective_type})',
                'list_level': item_level
            }
            elements.append(element_info)
            continue
        else:
            # End current list if it was ongoing
            if current_list:
                _update_list_end_para(elements, para_idx - 1)
                current_list = None

        # Detect illustrations (embedded images in paragraphs)
        if _has_embedded_image(paragraph):
            # Use the actual paragraph text as content
            illustration_content = text.strip() if text.strip() else 'Иллюстрация'
            # Use the correct GRAPHIC identifier based on counter
            graphic_ident = f"{effective_prefix}-GRAPHIC{illustration_counter}"
            element_info = {
                'type': 'illustration',
                'start_line': para_start_line,
                'start_char': para_start_char,
                'end_line': para_end_line,
                'end_char': para_end_char,
                'start_para': para_idx,
                'end_para': para_idx,
                'content': illustration_content,
                'xml_example': f'<figure><title>Название иллюстрации</title><graphic infoEntityIdent="{graphic_ident}"/></figure>',
                'details': 'Встраиваемая иллюстрация'
            }
            elements.append(element_info)
            illustration_counter += 1
            continue

        # Detect warnings/cautions
        if _is_warning(text):
            element_info = {
                'type': 'warning',
                'start_line': para_start_line,
                'start_char': para_start_char,
                'end_line': para_end_line,
                'end_char': para_end_char,
                'start_para': para_idx,
                'end_para': para_idx,
                'content': text,
                'xml_example': '<warning><warningAndCautionPara>Предупреждающий текст</warningAndCautionPara></warning>',
                'details': 'Предупреждение/Предупреждение'
            }
            elements.append(element_info)
            continue

        # Detect notes (text-based: "Примечание:", "ПРИМЕЧАНИЕ." etc.)
        if _is_note(text):
            element_info = {
                'type': 'note',
                'start_line': para_start_line,
                'start_char': para_start_char,
                'end_line': para_end_line,
                'end_char': para_end_char,
                'start_para': para_idx,
                'end_para': para_idx,
                'content': text,
                'xml_example': '<note><notePara>' + text + '</notePara></note>',
                'details': 'Примечание'
            }
            elements.append(element_info)
            continue

        # Detect references
        references = _find_references(text)
        for ref_data in references:
            if len(ref_data) == 4:  # illustration_reference with icn_ref
                ref_type, ref_number, context, icn_ref = ref_data
            else:  # other references
                ref_type, ref_number, context = ref_data
                icn_ref = None

            xml_example = {
                'table_reference': '<para>Ссылка на таблицу: <tableRef refType="tableref" refIdent="TAB0001"/></para>',
                'illustration_reference': f'<para>Ссылка на рисунок: <internalRef internalRefId="ICN{int(ref_number):02d}" internalRefTargetType="irtt01"/></para>',
                'data_module_reference': '<para>Ссылка на модуль данных: <dmRef refType="refdm" refIdent="DMC-S5-A-120-10-00-00A-011A-A"/></para>'
            }.get(ref_type, '<para>Ссылка</para>')

            # Add file info for illustration references
            if ref_type == 'illustration_reference':
                # Use sequential numbering for references
                if not hasattr(analyze_document_elements, 'global_illustration_ref_counter'):
                    analyze_document_elements.global_illustration_ref_counter = 0
                ref_number = analyze_document_elements.global_illustration_ref_counter + 1
                analyze_document_elements.global_illustration_ref_counter += 1

                # Use sequential numbering for graphic files (will be updated in post-processing)
                graphic_num = int(ref_number) - 1  # Will be corrected in post-processing
                graphic_file = f"{effective_prefix}-GRAPHIC{graphic_num}.jpg"
                details = f'Ссылка на {ref_type} {ref_number}, file: {graphic_file}'
            else:
                details = f'Ссылка на {ref_type} {ref_number}'

            element_info = {
                'type': ref_type,
                'start_line': para_start_line,
                'start_char': para_start_char,
                'end_line': para_end_line,
                'end_char': para_end_char,
                'start_para': para_idx,
                'end_para': para_idx,
                'content': f'Ссылка на {context} {ref_number}',
                'xml_example': xml_example,
                'details': details
            }
            elements.append(element_info)

        # Default: paragraph
        # If text has multiple lines, treat first line as title, rest as content
        if '\n' in text:
            lines = text.split('\n', 1)
            title = lines[0].strip()
            body = lines[1].strip()
            content = f'{title} {body}'
            xml_example = f'<para>{title} {body}</para>'
        else:
            content = text
            xml_example = f'<para>{text}</para>'

        element_info = {
            'type': 'paragraph',
            'start_line': para_start_line,
            'start_char': para_start_char,
            'end_line': para_end_line,
            'end_char': para_end_char,
            'start_para': para_idx,
            'end_para': para_idx,
            'content': content,
            'xml_example': xml_example,
            'details': 'Параграф'
        }
        elements.append(element_info)

    # Close any remaining list
    if current_list:
        _update_list_end_para(elements, len(doc.paragraphs) - 1)

    # Post-process elements to clean up table references in paragraphs that precede tables
    # and combine illustration references with figure name paragraphs
    processed_elements = []
    i = 0
    while i < len(elements):
        # Check for illustration_reference followed by figure name paragraph
        if (i < len(elements) - 1 and
            elements[i].get('type') == 'illustration_reference' and
            elements[i + 1].get('type') == 'paragraph'):
            figure_content = elements[i + 1].get('content', '').strip()
            figure_name_pattern = r'^[Рр]исунок\s*\d+\s*[–-]\s*.+'
            if re.match(figure_name_pattern, figure_content):
                # Extract ICN number from illustration_reference details
                details = elements[i].get('details', '')
                icn_match = re.search(r'illustration_reference (\d+)', details)
                icn_num = icn_match.group(1) if icn_match else '01'
                icn_ref = f"ICN{int(icn_num):02d}"

                # Use sequential graphic numbering starting from 0
                # This matches what extract_illustrations would save
                if not hasattr(analyze_document_elements, 'graphic_counter'):
                    analyze_document_elements.graphic_counter = 0
                graphic_num = analyze_document_elements.graphic_counter
                analyze_document_elements.graphic_counter += 1

                graphic_file = f"{effective_prefix}-GRAPHIC{graphic_num}.jpg"

                # Create combined illustration element
                # Use position information from illustration_positions if available
                ref_name = f"{effective_prefix}-GRAPHIC{graphic_num}"
                pos_info = illustration_positions.get(ref_name, {}) if illustration_positions else {}
                combined_elem = {
                    'type': 'illustration',
                    'start_line': pos_info.get('start_line', elements[i].get('start_line', 0)),
                    'start_char': pos_info.get('start_char', elements[i].get('start_char', 0)),
                    'end_line': pos_info.get('end_line', elements[i + 1].get('end_line', elements[i].get('end_line', 0))),
                    'end_char': pos_info.get('end_char', elements[i + 1].get('end_char', elements[i].get('end_char', 0))),
                    'start_para': pos_info.get('start_para', elements[i].get('start_para', 0)),
                    'end_para': pos_info.get('end_para', elements[i + 1].get('end_para', elements[i].get('end_para', 0))),
                    'content': figure_content,
                    'xml_example': f'<figure id="{icn_ref}"><title>{figure_content}</title><graphic infoEntityIdent="{effective_prefix}-GRAPHIC{graphic_num}" reproductionScale="32" reproductionWidth="170mm" reproductionHeight="120mm" id="g{int(icn_num)}"/></figure>',
                    'details': f'Иллюстрация {icn_num}, file: {graphic_file}',
                    'context_text': pos_info.get('context_text', '')
                }
                processed_elements.append(combined_elem)
                i += 2  # Skip both elements
                continue

        # Check for table references in paragraphs
        elem = elements[i]
        elem_type = elem.get('type', 'paragraph')
        content = elem.get('content', '')

        if elem_type == 'paragraph' and i < len(elements) - 1:
            next_elem = elements[i + 1]
            if next_elem.get('type') == 'table':
                table_title = next_elem.get('content', '')
                if table_title:
                    # Remove table references from paragraph content
                    patterns = [
                        r'\s*\([Тт]аблица\s*\d+\)\s*',
                        r'\s*[Тт]аблица\s*\d+\s*',
                        r'\s*[Тт]аб\.\s*\d+\s*',
                        r'\s*[Тт]абл\.\s*\d+\s*'
                    ]
                    for pattern in patterns:
                        content = re.sub(pattern, '', content, flags=re.IGNORECASE)

                    # Clean up extra whitespace and trailing punctuation
                    content = content.strip()
                    if content.endswith('.,'):
                        content = content[:-2]
                    elif content.endswith(','):
                        content = content[:-1]
                    elif content.endswith('.'):
                        pass  # Keep the period
                    else:
                        pass

                    # Update the element content and XML example
                    elem = elem.copy()
                    elem['content'] = content
                    elem['xml_example'] = f'<para>{content}</para>'

        processed_elements.append(elem)
        i += 1

    # Assign stable_id to every element based on final position, type, and content
    for idx, elem in enumerate(processed_elements):
        elem['stable_id'] = compute_stable_id(
            idx, elem.get('type', ''), elem.get('content', '')
        )

    return processed_elements


def _is_table_start(paragraph) -> bool:
    """Check if paragraph starts a table (simplified check)."""
    # This is a basic check; tables in docx don't directly correspond to paragraphs this way
    # In python-docx, tables are separate, but for logging purposes, we can detect them elsewhere
    # For now, return False; tables will be handled separately if needed
    return False


def _is_likely_list_item(text: str, paragraphs, para_idx: int, elements: list) -> bool:
    """
    Check if text is likely a list item rather than a header.
    This function is called BEFORE header detection to prevent false positives.

    Returns True if the text appears to be a list item based on:
    1. Starts with a list marker (bullet, dash, etc.)
    2. Ends with semicolon (typical list item ending)
    3. Previous paragraph ends with colon (list introduction)
    4. Previous element was also classified as a list item with similar structure
    5. Matches common list item patterns (e.g., "N шт. - description")
    """
    if not text:
        return False

    text = text.strip()

    # Check 0: Common Russian technical document list patterns
    # Pattern like "2 шт. - описание" (quantity + dash + description)
    quantity_pattern = r'^\d+\s*шт\.?\s*[-–—]'
    if re.match(quantity_pattern, text):
        return True

    # Check 0b: Numbered text with sentence case = list item, not header
    # e.g. "1. Описание процесса", "3.1.2 При монтаже необходимо"
    if _is_numbered_list_start(text):
        return True

    # Check 1: Explicit list markers at the start
    list_markers = ['•', '◦', '▪', '▫', '⁃', '·', '−', '–', '—', '†', '‡', '§']
    if any(text.startswith(marker) for marker in list_markers):
        return True

    # Check 1b: Dash at start (but not if it's a long text that looks like a heading)
    if text.startswith('-') or text.startswith('*'):
        # Only treat as list if it's followed by space and not all caps (which might be a header)
        if len(text) > 1 and text[1] == ' ' and not text.isupper():
            return True

    # Check 2: Ends with semicolon - strong indicator of list item
    if text.endswith(';'):
        return True

    # Check 3: Previous paragraph ends with colon (list introduction pattern)
    if para_idx > 0 and para_idx < len(paragraphs):
        prev_para = paragraphs[para_idx - 1]
        prev_text = prev_para.text.strip()
        if prev_text.endswith(':'):
            # Current paragraph follows a colon-intro, likely a list item
            # Unless it's clearly a section header (all caps, very short)
            if not (len(text) < 30 and text.isupper()):
                return True

    # Check 4: Previous element was a list item with similar structure
    # But skip if text looks like a numbered section header (ALL CAPS with number)
    if elements and not _is_main_section_header(text):
        prev_element = elements[-1]
        prev_type = prev_element.get('type', '')
        if prev_type in ('unnumbered_list', 'numbered_list'):
            prev_content = prev_element.get('content', '')
            # If previous was list and this has similar ending pattern
            if prev_content.endswith(';') and (text.endswith(';') or text.endswith('.')):
                return True
            # If previous was list and both are short without sentence endings
            if len(prev_content) < 150 and len(text) < 150:
                if not prev_content.endswith('.') and not text.endswith('.'):
                    return True

    # Check 5: Text looks like a continuation of a list (short, no period, similar to prev)
    if elements and len(text) < 100:
        prev_element = elements[-1]
        prev_type = prev_element.get('type', '')
        prev_content = prev_element.get('content', '')

        # If previous element ended with semicolon and this one does too
        if prev_type in ('unnumbered_list', 'numbered_list', 'numbered_paragraph_header'):
            if prev_content.endswith(';') and text.endswith(';'):
                return True

    return False


def _parse_punkt_style(style_name: str):
    """Detect Word styles like 'ПУНКТ 1', 'ПУНКТ 1.2', 'ПУНКТ 1.2.3'.

    These are custom heading styles in Russian technical documents where
    the numbering is auto-generated by Word (not present in paragraph text).

    Style naming convention:
      '01. ПУНКТ 1'     → level 1 header
      '08. ПУНКТ 1.2'   → level 2 header
      '16. ПУНКТ 1.2.3' → level 3 header

    Returns (level, pattern) tuple if matched, or None.
    Level is the nesting depth (1, 2, 3, ...).
    Pattern is the numbering template (e.g. '1.2.3').
    """
    if not style_name:
        return None
    # Match styles containing 'ПУНКТ' (case-insensitive) followed by a
    # numbering pattern like '1', '1.2', '1.2.3'
    m = re.search(r'(?i)ПУНКТ\s+(\d+(?:\.\d+)*)', style_name)
    if not m:
        return None
    pattern = m.group(1)
    level = pattern.count('.') + 1
    return (level, pattern)


# Counters for ПУНКТ-style numbering, keyed by level.
# Reset for each document analysis run.
_punkt_counters = {}


def _get_punkt_number(level: int) -> str:
    """Generate hierarchical number for ПУНКТ-style headers.

    Maintains per-level counters. When a higher-level header appears,
    resets all lower-level counters.

    Example sequence:
      level 1 → "1"
      level 2 → "1.1"
      level 2 → "1.2"
      level 1 → "2"
      level 2 → "2.1"
      level 3 → "2.1.1"
    """
    global _punkt_counters

    # Increment counter at this level
    _punkt_counters[level] = _punkt_counters.get(level, 0) + 1

    # Reset all deeper levels
    for lvl in list(_punkt_counters.keys()):
        if lvl > level:
            _punkt_counters[lvl] = 0

    # Build hierarchical number: 1, 1.1, 1.1.1, etc.
    parts = []
    for lvl in range(1, level + 1):
        parts.append(str(_punkt_counters.get(lvl, 0)))
    return '.'.join(parts)


def _is_main_section_header(text: str) -> bool:
    """Check if text is a section header: number(s) + ALL CAPS text.

    Matches: '1 МЕРЫ БЕЗОПАСНОСТИ', '3.1.2 ОПИСАНИЕ', '2. ОБЩИЕ СВЕДЕНИЯ'
    Does NOT match: '1. Описание процесса', '3.1.2 При монтаже'
    """
    pattern = r'^\s*\d+(?:\.\d+)*\.?\s+'
    m = re.match(pattern, text)
    if not m:
        return False
    remainder = text[m.end():].strip()
    if not remainder:
        return False
    letters = [c for c in remainder if c.isalpha()]
    if not letters:
        return False
    upper_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
    return upper_ratio >= 0.7


def _is_numbered_list_start(text: str) -> bool:
    """Check if text is a numbered list item with sentence-case text.

    Matches: '1. Описание', '3.1.2 При монтаже', '2) Порядок'
    Does NOT match: '1 МЕРЫ БЕЗОПАСНОСТИ', '3.1.2 ОПИСАНИЕ', '3 шт. - описание'

    Single numbers require a dot or paren separator (1., 2)).
    Multi-level numbers (3.1.2) don't need an extra separator.
    """
    # Multi-level numbers (3.1, 3.1.2) with optional trailing dot
    # OR single number with explicit separator (dot or paren)
    pattern = r'^\s*\d+(?:(?:\.\d+)+\.?|[\.\)])\s+'
    m = re.match(pattern, text)
    if not m:
        return False
    remainder = text[m.end():].strip()
    if not remainder:
        return False
    letters = [c for c in remainder if c.isalpha()]
    if not letters:
        return False
    upper_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
    return upper_ratio < 0.7


def _get_list_level(paragraph) -> int:
    """Get list nesting level from DOCX.

    Tries in order: numPr.ilvl, raw XML ilvl, style left_indent,
    paragraph left_indent. Returns a numeric value where higher = more nested.
    The processor normalizes these to relative 0/1/2... levels.
    """
    try:
        numPr = getattr(paragraph.paragraph_format, 'numPr', None)
        if numPr is not None:
            ilvl = getattr(numPr, 'ilvl', None)
            if ilvl is not None and hasattr(ilvl, 'val'):
                return ilvl.val
        # Fallback: check raw XML for ilvl
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        pPr = paragraph._element.find('.//w:pPr', ns)
        if pPr is not None:
            numPr_elem = pPr.find('.//w:numPr', ns)
            if numPr_elem is not None:
                ilvl_elem = numPr_elem.find('.//w:ilvl', ns)
                if ilvl_elem is not None:
                    return int(ilvl_elem.get(f'{{{ns["w"]}}}val', '0'))
    except Exception:
        pass

    # Fallback: use effective left indent (paragraph-level, then style-level).
    # Returns raw EMU value — the processor normalizes relative levels.
    try:
        li = paragraph.paragraph_format.left_indent
        if li is not None and li > 0:
            return li
        if paragraph.style:
            style_li = paragraph.style.paragraph_format.left_indent
            if style_li is not None and style_li > 0:
                return style_li
    except Exception:
        pass

    return 0


def _get_list_type(paragraph, para_idx: int = None) -> str:
    """Get list type for paragraph based on OOXML formatting and enhanced heuristics."""
    text = paragraph.text.strip()

    # Skip section headers (ALL CAPS with numbers)
    if _is_main_section_header(text):
        return ''

    # Sentence-case numbered items are numbered lists
    if _is_numbered_list_start(text):
        return 'numbered_list'

    # First try OOXML formatting (preferred method)
    if hasattr(paragraph.paragraph_format, 'numPr') and paragraph.paragraph_format.numPr is not None:
        # Get numId to determine list type
        num_id = paragraph.paragraph_format.numPr.numId
        if num_id is not None:
            # Access document numbering to determine if numbered or bulleted
            doc = paragraph.part.document
            numbering_part = doc.numbering_part
            if numbering_part is not None:
                try:
                    num = numbering_part.nums[num_id.val]
                    abstract_num = numbering_part.abstract_nums[num.abstract_num_id.val]
                    # Check lvl0 lvlText to determine list type - if contains % it's numbered
                    lvl0 = abstract_num.lvl0
                    if lvl0 is not None and hasattr(lvl0, 'lvlText') and lvl0.lvlText is not None:
                        lvl_text = lvl0.lvlText.val
                        if '%' in lvl_text:
                            return 'numbered_list'
                        else:
                            return 'unnumbered_list'
                    else:
                        # Fallback to numFmt if lvlText not available
                        if hasattr(abstract_num, 'num_style_link') and abstract_num.num_style_link is not None:
                            # If linked to a style, check if it's a numbered style
                            style_name = abstract_num.num_style_link.val.lower()
                            if 'number' in style_name or 'decimal' in style_name:
                                return 'numbered_list'
                            else:
                                return 'unnumbered_list'
                        else:
                            # Check lvl0 numFmt directly
                            lvl0 = abstract_num.lvl0
                            if lvl0 is not None and hasattr(lvl0, 'numFmt') and lvl0.numFmt is not None:
                                num_fmt = lvl0.numFmt.val.lower()
                                if num_fmt in ['decimal', 'lowerroman', 'upperroman', 'lowerletter', 'upperletter']:
                                    return 'numbered_list'
                                else:
                                    return 'unnumbered_list'
                            else:
                                return 'unnumbered_list'
                except (KeyError, AttributeError):
                    return 'unnumbered_list'
            else:
                return 'unnumbered_list'
        else:
            return 'unnumbered_list'

    # Fallback to enhanced heuristics if no OOXML formatting found
    doc = paragraph.part.document
    
    # If para_idx not provided, find it
    if para_idx is None:
        for idx, p in enumerate(doc.paragraphs):
            if p == paragraph:
                para_idx = idx
                break
    
    if para_idx is None:
        return ''
    
    # Check for list markers in text
    unnumbered_markers = ['•', '◦', '▪', '▫', '⁃', '·', '-', '–', '—', '*', '†', '‡', '§']
    if any(text.startswith(marker) for marker in unnumbered_markers):
        return 'unnumbered_list'
    
    # Check if text ends with ';' (typical for list items)
    if text.endswith(';'):
        # Strong indicator of list item - text ending with semicolon
        # Either verify it's after an intro paragraph, or if text is short enough to be a list item
        if _is_after_colon_intro(doc, para_idx):
            return 'unnumbered_list'
        # PHASE 1 FIX: Also treat as list item if it looks like a list item
        # Text ending with ; is very likely a list item - just verify it's not too long
        # and doesn't look like a complete multi-sentence paragraph
        if len(text) < 300 and text.count('. ') <= 2:
            return 'unnumbered_list'

    # Check colon-intro pattern
    list_type = _detect_colon_intro_list(doc, para_idx)
    if list_type:
        return list_type

    return ''


def _is_after_colon_intro(doc, para_idx: int) -> bool:
    """Check if paragraph is after an introductory paragraph ending with ':'."""
    paragraphs = doc.paragraphs
    
    # Search backward for introductory paragraph ending with ":"
    for idx in range(para_idx - 1, -1, -1):
        para = paragraphs[idx]
        text = para.text.strip()
        if not text:
            # Empty line - check if it breaks the list
            continue
        if text.endswith(':'):
            return True
        # If we hit a paragraph that doesn't end with ':', ';' or is empty, stop
        if not text.endswith(';') and not text.endswith(':'):
            break
    
    return False


def _detect_colon_intro_list(doc, para_idx: int) -> str:
    """
    Detect if paragraph is part of a list that starts after a colon-ending intro.
    """
    paragraphs = doc.paragraphs
    text = paragraphs[para_idx].text.strip()
    
    # Search backward for introductory paragraph ending with ":"
    intro_idx = None
    for idx in range(para_idx - 1, -1, -1):
        para = paragraphs[idx]
        para_text = para.text.strip()
        if not para_text:
            # Empty line means we've passed the list boundary
            break
        if para_text.endswith(':'):
            intro_idx = idx
            break
    
    if intro_idx is None:
        return ''
    
    # Check if current paragraph is within the list items after intro
    # List ends at empty line or at another line ending with ':'
    current_idx = intro_idx + 1
    list_items_indices = []
    
    while current_idx < len(paragraphs):
        para = paragraphs[current_idx]
        para_text = para.text.strip()
        
        # Empty line ends the list
        if not para_text:
            break
        
        # Another introductory line ends this list
        if para_text.endswith(':'):
            break
        
        list_items_indices.append(current_idx)
        current_idx += 1
    
    # If current paragraph is in the list items, return unnumbered_list
    if para_idx in list_items_indices and len(list_items_indices) >= 1:
        return 'unnumbered_list'
    
    return ''


def _clean_list_item_text(paragraph, list_type: str) -> str:
    """Clean list item text by removing markers and formatting."""
    text = paragraph.text.strip()

    if list_type == 'unnumbered_list':
        # Remove common unnumbered list markers
        unnumbered_markers = [
            '•', '◦', '▪', '▫', '⁃', '·', '·', '•', '-', '–', '—', '*', '†', '‡', '§'
        ]
        for marker in unnumbered_markers:
            if text.startswith(marker):
                text = text[len(marker):].strip()
                break
    elif list_type == 'numbered_list':
        # Remove numbered list markers using regex
        import re
        numbered_patterns = [
            r'^\d+\.\s*',  # 1., 2., 10.
            r'^\d+\)\s*',  # 1), 2), 10)
            r'^\(\d+\)\s*',  # (1), (2), (10)
            r'^\d+\s*[.)]\s*',  # 1 , 1), 2 , etc.
            r'^[a-zA-Z]\.\s*',  # a., b., A., B.
            r'^[a-zA-Z]\)\s*',  # a), b), A), B)
            r'^\([a-zA-Z]\)\s*',  # (a), (b), (A), (B)
            r'^[IVXLCDM]+\.\s*',  # I., II., III. (Roman numerals)
            r'^[ivxlcdm]+\.\s*',  # i., ii., iii. (Roman numerals)
            r'^[IVXLCDM]+\)\s*',  # I), II), III)
            r'^[ivxlcdm]+\)\s*',  # i), ii), iii)
        ]
        for pattern in numbered_patterns:
            text = re.sub(pattern, '', text, count=1)
            if text != paragraph.text.strip():  # If something was removed
                break

    return text.strip()


def _update_list_end_para(elements: List[Dict], end_para: int):
    """Update end_para for the last list items."""
    for element in reversed(elements):
        if element['type'] in ['numbered_list', 'unnumbered_list']:
            element['end_para'] = end_para
            if element.get('xml_example') and 'listItem' in element['xml_example']:
                # It's a single item, no need to update further
                break
        else:
            break


def _has_embedded_image(paragraph) -> bool:
    """Check if paragraph contains embedded image."""
    for run in paragraph.runs:
        if run.element.xpath('.//a:blip'):
            return True
    return False


def _is_warning(text: str) -> bool:
    """Check if text contains warning/caution content.

    Uses word-boundary matching to avoid false positives like
    'безОПАСНОсти' matching 'опасно'.
    """
    lower_text = text.lower()
    warning_keywords = ['внимание', 'осторожно', 'предупреждение', 'caution', 'warning', 'опасно']
    for keyword in warning_keywords:
        if re.search(r'\b' + keyword + r'\b', lower_text):
            return True
    return False


def _is_warning_style(style_name: str) -> bool:
    """Detect warning styles by name: '07. Внимание для п.1', '22. Внимание для п.1.2.3'."""
    if not style_name:
        return False
    return bool(re.search(r'(?i)вниман', style_name))


def _is_note_style(style_name: str) -> bool:
    """Detect note/remark styles by name: '07. Замеч для п.1', '07. Примеч. для 1.2'."""
    if not style_name:
        return False
    return bool(re.search(r'(?i)(примеч|замеч)', style_name))


def _is_note(text: str) -> bool:
    """Detect note by text content: 'Примечание:', 'ПРИМЕЧАНИЕ.', etc."""
    return bool(re.match(r'(?i)^\s*примечани[ея][\s.:–\-]', text))


def _find_references(text: str) -> List[Tuple[str, str, str]]:
    """Find various references in text."""
    references = []

    # Table references
    table_patterns = [
        r'\b[Тт]аблица\s*(\d+)',
        r'\b[Тт]абл\.\s*(\d+)',
        r'\bTable\s*(\d+)',
    ]
    for pattern in table_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            references.append(('table_reference', match, 'таблицу'))

    # Illustration/Figure references - include regular figure mentions
    figure_patterns = [
        r'^[Рр]исунок\s*(\d+)\s*[–-]\s*.+',
        r'\b[Рр]ис\.\s*(\d+)',
        r'\b[Фф]igure\s*(\d+)',
    ]
    for pattern in figure_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            # ICN will be assigned sequentially later in processing
            references.append(('illustration_reference', match, 'иллюстрацию'))

    # Data module references (more complex, might be DM codes)
    dm_patterns = [
        r'\bDMC-[A-Z]\d+-[A-Z]-(\d+)-.*?\b',
        r'\bМодуль\s*данных\s*(\d+)',
    ]
    for pattern in dm_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            for match in matches:
                references.append(('data_module_reference', match, 'модуль данных'))

    return references


def apply_overrides(elements: List[Dict[str, Any]], dmc_string: str) -> List[Dict[str, Any]]:
    """
    Apply verification overrides to classified elements.

    Overrides are generated by verify_loop.py after comparing the current XML
    output against a user-edited reference markup.

    Args:
        elements: List of element dicts from analyze_document_elements()
        dmc_string: DMC identifier to look up overrides

    Returns:
        Modified elements list with overrides applied.
    """
    from app_paths import get_app_root
    overrides_dir = os.path.join(
        get_app_root(),
        'comparison_app', '_overrides',
    )
    overrides_path = os.path.join(overrides_dir, f'{dmc_string}.json')

    if not os.path.isfile(overrides_path):
        return elements

    with open(overrides_path, 'r', encoding='utf-8') as f:
        overrides = json.load(f)

    # Type mapping from reference types to elements_analyzer types
    type_map = {
        'heading': 'numbered_paragraph_header',
        'para': 'paragraph',
        'numbered_list': 'numbered_list',
        'unnumbered_list': 'unnumbered_list',
        'list': 'unnumbered_list',  # backward compat
        'table': 'table',
        'figure': 'illustration',
        'warning': 'warning',
        'caution': 'warning',
        'note': 'paragraph',
    }

    # --- Text-based reclassify rules (new format) ---
    reclassify_rules = overrides.get('reclassify_rules', [])
    for rule in reclassify_rules:
        target_text = rule.get('text_start', '')
        target_end = rule.get('text_end', '')
        new_type = rule.get('to_type', '')
        mapped_type = type_map.get(new_type, new_type)
        if not target_text or not mapped_type:
            continue
        for elem in elements:
            content = elem.get('content', '')
            content_start = content[:60].strip()
            content_end = content[-40:].strip() if len(content) > 40 else content.strip()
            # Match by text_start prefix (primary) + optional text_end check
            if content_start and target_text.strip() and content_start.startswith(target_text.strip()[:30]):
                if not target_end or content_end.endswith(target_end.strip()[-20:]):
                    elem['type'] = mapped_type
                    break

    # --- Text-based skip rules (new format) ---
    skip_rules = overrides.get('skip_rules', [])
    if skip_rules:
        skip_texts = set()
        for rule in skip_rules:
            text = rule.get('text_start', '').strip()
            if text:
                skip_texts.add(text[:30])  # Use first 30 chars as key

        def _should_skip(elem):
            content = elem.get('content', '')[:30].strip()
            return content in skip_texts if content else False

        elements = [e for e in elements if not _should_skip(e)]

    # --- Legacy index-based overrides (backward compat) ---
    reclassify_legacy = overrides.get('reclassify', {})
    skip_legacy = set(overrides.get('skip_elements', []))

    # Only apply legacy if no new-format rules exist
    if not reclassify_rules and reclassify_legacy:
        # Legacy: keys were XML idx, try matching by list position (approximate)
        for str_idx, new_type in reclassify_legacy.items():
            try:
                idx = int(str_idx)
                # Offset by -1 because XML extraction adds a heading element at idx=1
                elem_idx = idx - 1
                if 0 <= elem_idx < len(elements):
                    mapped_type = type_map.get(new_type, new_type)
                    elements[elem_idx]['type'] = mapped_type
            except (ValueError, IndexError):
                pass

    if not skip_rules and skip_legacy:
        # Legacy: values were XML idx (1-based), convert to 0-based
        adjusted = set()
        for idx in skip_legacy:
            adjusted.add(idx - 1)  # Convert from 1-based to 0-based
        elements = [e for i, e in enumerate(elements) if i not in adjusted]

    return elements


# Module-level cache for last markup result (used by verify_loop for XSD mapping)
_last_markup_result: Dict[str, List[Dict[str, Any]]] = {}


def get_last_markup_result(dmc_string: str) -> Optional[List[Dict[str, Any]]]:
    """Return cached elements from last apply_reference_markup() call for XSD mapping."""
    return _last_markup_result.get(dmc_string)


def apply_reference_markup(elements: List[Dict[str, Any]], dmc_string: str) -> List[Dict[str, Any]]:
    """
    Apply user reference markup directly to classified elements.

    For each reference element, finds the matching auto-extracted element by text
    similarity and overrides its type IN PLACE (no reordering). Elements not
    covered by reference keep their auto-classified type.

    Each element is tagged with:
      - _ref_annotated: True if type was set from user reference, False otherwise
      - _original_type: the auto-classified type before override (if overridden)
      - _ref_idx: reference element index (if matched)
      - _ref_type_raw: raw reference type string (if matched)
      - _element_id: stable element identifier from hybrid pipeline (if present in ref)

    Falls back to apply_overrides() if no reference markup exists.
    """
    # Try to load reference markup
    try:
        from comparison_app.reference_store import get_reference
        ref_data = get_reference(dmc_string)
    except ImportError:
        ref_data = None

    if not ref_data or not ref_data.get('elements'):
        # No reference — fall back to overrides only
        return apply_overrides(elements, dmc_string)

    ref_elements = ref_data['elements']

    # Filter out sentinel types from XML-first references — they have no
    # real DOCX counterpart and would steal matches from actual elements.
    ref_elements = [e for e in ref_elements
                    if e.get('type') not in ('_extra_pdf', '_unmatched_xml')]

    # Type mapping: reference types → analyzer types used by descriptive_processor
    type_map = {
        'heading': 'header',
        'para': 'paragraph',
        'paragraph': 'paragraph',
        'numbered_list': 'numbered_list',
        'unnumbered_list': 'unnumbered_list',
        'list': 'unnumbered_list',
        'table': 'table',
        'figure': 'illustration',
        'warning': 'warning',
        'caution': 'caution',
        'note': 'note',
    }

    def _prefix_score(ref_text, content_text):
        """Compute prefix-based similarity score (case-insensitive)."""
        if not ref_text or not content_text:
            return 0.0
        a = ref_text.strip().lower()
        b = content_text.strip().lower()
        min_len = min(len(a), len(b))
        if min_len == 0:
            return 0.0
        prefix_len = 0
        for i in range(min_len):
            if a[i] == b[i]:
                prefix_len += 1
            else:
                break
        if prefix_len < 3:
            return 0.0
        return prefix_len / max(len(a), 1)

    def _collapse_spaced_letters(text):
        """Collapse spaced-out letters back to words.

        PDF sometimes renders 'Примечание' as 'П р и м е ч а н и е'
        (each letter separated by space). Collapse these back to normal words.
        """
        import re
        # Match runs of: single-char + space + single-char + space ...
        # At least 3 letters with spaces between them
        def _collapse(m):
            return m.group(0).replace(' ', '')
        return re.sub(r'(?<!\S)\S(?:\s\S){2,}(?!\S)', _collapse, text)

    def _strip_list_prefix(text):
        """Strip leading numbering and list markers for better matching.

        PDF text often has section numbers ('1 TEXT', '3.1 TEXT') and
        dash bullets ('– text') that DOCX auto-extraction removes.
        """
        import re
        t = text.strip()
        # Strip leading section numbers: "1 TEXT" → "TEXT", "3.1.4 TEXT" → "TEXT"
        t = re.sub(r'^\d+(?:\.\d+)*[\.\)]*\s*', '', t)
        # Strip leading dashes/bullets: "– TEXT" → "TEXT", "- TEXT" → "TEXT"
        t = re.sub(r'^[–\-•]\s*', '', t)
        return t

    def _combined_score(ref_text_start, ref_text_end, content):
        """Score using both text_start and text_end for better disambiguation."""
        content_start = content[:60].strip()
        content_end = content[-40:].strip() if len(content) > 40 else content.strip()
        start_score = _prefix_score(ref_text_start, content_start)

        # Fallback: try with stripped list/number prefixes
        stripped_ref = _strip_list_prefix(ref_text_start)
        stripped_content = _strip_list_prefix(content_start)
        if stripped_ref and stripped_content:
            alt_score = _prefix_score(stripped_ref, stripped_content)
            start_score = max(start_score, alt_score)

        # Fallback: try with collapsed spaced-out letters
        # PDF renders "Примечание" as "П р и м е ч а н и е"
        collapsed_ref = _collapse_spaced_letters(ref_text_start)
        if collapsed_ref != ref_text_start:
            alt_score = _prefix_score(collapsed_ref, content_start)
            start_score = max(start_score, alt_score)
            # Also try collapsed + stripped
            stripped_collapsed = _strip_list_prefix(collapsed_ref)
            if stripped_collapsed:
                alt_score = _prefix_score(stripped_collapsed, stripped_content or content_start)
                start_score = max(start_score, alt_score)

        if not ref_text_end or not ref_text_end.strip():
            return start_score
        end_score = _prefix_score(ref_text_end, content_end)
        # Also try collapsed ref_text_end
        collapsed_end = _collapse_spaced_letters(ref_text_end)
        if collapsed_end != ref_text_end:
            end_score = max(end_score, _prefix_score(collapsed_end, content_end))
        return start_score * 0.7 + end_score * 0.3

    def _find_best_match(cursor, ref_text_start, ref_text_end, used):
        """Find best matching auto element by text similarity."""
        if not ref_text_start:
            return None

        min_threshold = 0.2 if len(ref_text_start) < 10 else 0.3
        best_idx = None
        best_score = min_threshold

        # Search forward from cursor (window: cursor-5 .. cursor+50)
        search_start = max(0, cursor - 5)
        search_end = min(len(elements), cursor + 50)

        for i in range(search_start, search_end):
            if i in used:
                continue
            content = elements[i].get('content', '')
            score = _combined_score(ref_text_start, ref_text_end, content)
            if score > best_score:
                best_score = score
                best_idx = i

        # Fallback: search entire list
        if best_idx is None:
            for i in range(len(elements)):
                if i in used:
                    continue
                content = elements[i].get('content', '')
                score = _combined_score(ref_text_start, ref_text_end, content)
                if score > best_score:
                    best_score = score
                    best_idx = i

        return best_idx

    def _apply_span_forward(start_idx, ref_span, ref_text_end, mapped_type, ref_idx, ref_type_raw, used):
        """For span>1 refs, scan forward from start_idx overriding types.

        Scans forward up to ref_span * 1.5 auto elements (to account for
        granularity difference between DOM blocks and auto elements) until:
        - An element containing ref_text_end is found (inclusive)
        - Or max scan limit is reached
        - Or an already-used element is encountered (except _skip elements,
          which are skipped over without consuming them)
        """
        max_scan = min(len(elements), start_idx + 1 + int(ref_span * 1.5))
        for i in range(start_idx + 1, max_scan):
            if i in used:
                # Skip over _skip elements — they don't block span_forward
                if elements[i].get('type') == '_skip':
                    continue
                break
            # Illustrations break text groups (lists, paragraphs) — never
            # override an illustration with a non-illustration type.
            elem_type = elements[i].get('type', '')
            if elem_type == 'illustration' and mapped_type != 'illustration':
                break
            elements[i]['_original_type'] = elements[i]['type']
            # Section headers: only the first element is the header;
            # subsequent span elements are content paragraphs within that section
            if mapped_type == 'numbered_paragraph_header':
                elements[i]['type'] = 'paragraph'
            else:
                elements[i]['type'] = mapped_type
            elements[i]['_ref_annotated'] = True
            elements[i]['_ref_idx'] = ref_idx
            elements[i]['_ref_type_raw'] = ref_type_raw
            used.add(i)

            # Stop if this element contains ref_text_end (with prefix normalization)
            if ref_text_end:
                content_end = elements[i].get('content', '')[-40:].strip()
                if content_end:
                    score = _prefix_score(ref_text_end, content_end)
                    # Also try with stripped list prefixes
                    stripped_ref_end = _strip_list_prefix(ref_text_end)
                    stripped_content_end = _strip_list_prefix(content_end)
                    if stripped_ref_end and stripped_content_end:
                        score = max(score, _prefix_score(stripped_ref_end, stripped_content_end))
                    if score > 0.3:
                        break

    # --- Main matching loop: modify elements IN PLACE, no reordering ---
    # Three-phase approach:
    #   1. Match each non-skip ref to its best single auto element (ignore span)
    #   2. Match _skip refs on remaining unused elements
    #   3. Apply deferred span_forward calls on remaining unused elements
    # This prevents span_forward from consuming elements that a later ref
    # could match individually (e.g., ref[36] span=4 would grab elements
    # that ref[37..40] should claim).
    used = set()
    cursor = 0
    skip_refs = []       # Deferred _skip processing
    span_deferred = []   # Deferred span_forward calls

    for ref_elem in ref_elements:
        ref_type = ref_elem.get('type', '')
        ref_text_start = (ref_elem.get('text_start', '') or '').strip()
        ref_text_end = (ref_elem.get('text_end', '') or '').strip()
        ref_span = ref_elem.get('span', 1) or 1
        ref_idx = ref_elem.get('idx', 0)

        # Defer _skip elements to second pass
        if ref_type == '_skip':
            skip_refs.append(ref_elem)
            continue

        mapped_type = type_map.get(ref_type, ref_type)

        # Section-numbered lists → numbered_paragraph_header ONLY for depth <= 2
        # ("3 Общие указания" depth=1, "3.2 Последовательность" depth=2 → <levelledPara>)
        # ("3.2.1 Снимите...", "3.1.4 При монтаже..." depth=3+ → keep as numbered_list)
        if mapped_type in ('numbered_list', 'nested_numbered_list') and ref_text_start:
            m_sec = re.match(r'^(\d+(?:\.\d+)*)\s+', ref_text_start)
            if m_sec and m_sec.group(1).count('.') < 2:
                mapped_type = 'numbered_paragraph_header'

        ref_element_id = ref_elem.get('element_id', '')

        best_idx = _find_best_match(cursor, ref_text_start, ref_text_end, used)

        if best_idx is not None:
            orig_type = elements[best_idx]['type']
            orig_content = elements[best_idx].get('content', '')[:60]
            print(f'[ref_markup] ref[{ref_idx}] {ref_type}->{mapped_type} '
                  f'matched auto[{best_idx}] (was {orig_type}) '
                  f'"{ref_text_start[:30]}" <-> "{orig_content[:30]}"')
            # Override type in place
            elements[best_idx]['_original_type'] = elements[best_idx]['type']
            elements[best_idx]['type'] = mapped_type
            elements[best_idx]['_ref_annotated'] = True
            elements[best_idx]['_ref_idx'] = ref_idx
            elements[best_idx]['_ref_type_raw'] = ref_type
            if ref_element_id:
                elements[best_idx]['_element_id'] = ref_element_id
            used.add(best_idx)
            cursor = best_idx + 1

            # Defer span_forward to phase 3 (after all single-element matching).
            # Always defer when span > 1 — span_forward handles the "already
            # complete" case gracefully (it will simply not find extra elements).
            # Removing the previous already_complete optimization fixes user merge
            # scenarios where span>1 covering multiple DOCX paragraphs.
            if ref_span > 1:
                span_deferred.append((best_idx, ref_span, ref_text_end,
                                     mapped_type, ref_idx, ref_type))
                # Advance cursor past expected span range so later refs
                # search beyond this span
                cursor = max(cursor, best_idx + ref_span)

    # --- Phase 2: process _skip elements on remaining unused auto elements ---
    for ref_elem in skip_refs:
        ref_text_start = (ref_elem.get('text_start', '') or '').strip()
        ref_text_end = (ref_elem.get('text_end', '') or '').strip()
        ref_idx = ref_elem.get('idx', 0)
        best_idx = _find_best_match(0, ref_text_start, ref_text_end, used)
        if best_idx is not None:
            orig_type = elements[best_idx]['type']
            orig_content = elements[best_idx].get('content', '')[:60]
            print(f'[ref_markup] _skip matched auto[{best_idx}] (was {orig_type}) '
                  f'"{ref_text_start[:30]}" <-> "{orig_content[:30]}"')
            elements[best_idx]['_original_type'] = elements[best_idx]['type']
            elements[best_idx]['type'] = '_skip'
            elements[best_idx]['_ref_annotated'] = True
            elements[best_idx]['_ref_idx'] = ref_idx
            elements[best_idx]['_ref_type_raw'] = '_skip'
            used.add(best_idx)

    # --- Phase 3: apply deferred span_forward on remaining unused elements ---
    for args in span_deferred:
        start_idx, ref_span_val = args[0], args[1]
        print(f'[ref_markup] Phase 3: span_forward from auto[{start_idx}], '
              f'span={ref_span_val}, type={args[3]}, ref_idx={args[4]}')
        _apply_span_forward(*args, used)

    # --- Phase 4: post-match fixups for granularity mismatches ---
    # Build ref_idx -> matched_auto_idx lookup
    matched_by_ref = {}  # ref_idx -> auto_idx
    for i, elem in enumerate(elements):
        if elem.get('_ref_annotated') and '_ref_idx' in elem:
            matched_by_ref[elem['_ref_idx']] = i

    # 4a. Extend: if a matched ref's text_end falls in an adjacent unused element,
    #     include that element in the match (PDF block covers multiple DOCX elements)
    def _ends_match(ref_end, content_tail):
        """Check if ref text_end matches content tail using prefix + suffix fallback."""
        if not ref_end or not content_tail:
            return False
        score = _prefix_score(ref_end, content_tail)
        sr = _strip_list_prefix(ref_end)
        sc = _strip_list_prefix(content_tail)
        if sr and sc:
            score = max(score, _prefix_score(sr, sc))
        if score > 0.3:
            return True
        # Fallback: suffix matching (content[-40:] truncation can shift prefix)
        norm_ref = ref_end.lower().replace('\n', ' ').replace('\r', ' ').strip()
        norm_ct = content_tail.lower().replace('\n', ' ').replace('\r', ' ').strip()
        suffix_len = min(15, len(norm_ref), len(norm_ct))
        if suffix_len >= 8 and norm_ref[-suffix_len:] == norm_ct[-suffix_len:]:
            return True
        return False

    for ref_elem in ref_elements:
        ref_type = ref_elem.get('type', '')
        if ref_type == '_skip':
            continue
        ref_text_end = (ref_elem.get('text_end', '') or '').strip()
        ref_idx = ref_elem.get('idx', 0)
        if not ref_text_end:
            continue

        matched_idx = matched_by_ref.get(ref_idx)
        if matched_idx is None:
            continue

        # Check if matched element's content already covers text_end
        content_end = elements[matched_idx].get('content', '')[-40:].strip()
        if _ends_match(ref_text_end, content_end):
            continue  # Already complete

        # text_end not covered — check adjacent unused elements
        mapped_type = elements[matched_idx]['type']
        for j in range(matched_idx + 1, min(len(elements), matched_idx + 5)):
            if j in used:
                break  # Stop at used element
            next_content_end = elements[j].get('content', '')[-40:].strip()
            if not next_content_end:
                continue
            if _ends_match(ref_text_end, next_content_end):
                # Extend: include elements [matched_idx+1 .. j] in this match
                for k in range(matched_idx + 1, j + 1):
                    if k not in used:
                        elements[k]['_original_type'] = elements[k]['type']
                        elements[k]['type'] = mapped_type
                        elements[k]['_ref_annotated'] = True
                        elements[k]['_ref_idx'] = ref_idx
                        elements[k]['_ref_type_raw'] = ref_type
                        used.add(k)
                        print(f'[ref_markup] extend ref[{ref_idx}] -> auto[{k}] '
                              f'(was {elements[k]["_original_type"]})')
                break

    # 4b. Split: if an unmatched ref's text_start is found WITHIN a matched element's
    #     content, add split metadata so the processor generates separate <para> elements
    #     (one DOCX paragraph contains text from multiple PDF blocks)
    for ref_elem in ref_elements:
        ref_type = ref_elem.get('type', '')
        if ref_type == '_skip':
            continue
        ref_idx = ref_elem.get('idx', 0)
        ref_text_start = (ref_elem.get('text_start', '') or '').strip()
        if not ref_text_start or ref_idx in matched_by_ref:
            continue  # Already matched

        # Search matched elements for one containing this text
        search_prefix = ref_text_start[:30].lower()
        for i, elem in enumerate(elements):
            if not elem.get('_ref_annotated') or elem.get('type') == '_skip':
                continue
            content = elem.get('content', '')
            content_lower = content.lower()
            pos = content_lower.find(search_prefix)
            if pos > 5:  # Found as substring, not at the start
                mapped_type = type_map.get(ref_type, ref_type)
                if '_split_points' not in elem:
                    elem['_split_points'] = []
                elem['_split_points'].append({
                    'position': pos,
                    'type': mapped_type,
                    'ref_idx': ref_idx,
                    'ref_type_raw': ref_type,
                })
                matched_by_ref[ref_idx] = i  # Mark as handled
                print(f'[ref_markup] split auto[{i}] at pos {pos} for ref[{ref_idx}] '
                      f'{ref_type} "{ref_text_start[:30]}"')
                break

    # Mark unannotated elements
    for i, elem in enumerate(elements):
        if i not in used:
            elem['_ref_annotated'] = False

    matched = len(used)
    unmatched = len(elements) - matched
    print(f'[apply_reference_markup] {matched}/{len(elements)} auto elements matched '
          f'by reference ({len(ref_elements)} ref elements, {unmatched} uncovered)')

    # Cache for XSD error mapping
    _last_markup_result[dmc_string] = elements

    # Document order preserved — return elements as-is
    return elements


def generate_elements_log(document_path: str, elements: List[Dict[str, Any]], output_path: str) -> str:
    """
    Generate elements log file for a document.

    Args:
        document_path: Path to source document
        elements: List of element dictionaries
        output_path: Directory to save log file

    Returns:
        Path to generated log file
    """
    filename = os.path.basename(document_path).replace('.docx', '').replace('.doc', '')
    _logs_dir = os.path.join(output_path, '_logs')
    os.makedirs(_logs_dir, exist_ok=True)
    log_path = os.path.join(_logs_dir, f"elements_{filename}.log")

    element_type_names = {
        'header': 'Заголовок',
        'numbered_paragraph_header': 'Нумерованный заголовок параграфа',
        'paragraph': 'Параграф',
        'numbered_list': 'Нумерованный список',
        'unnumbered_list': 'Ненумерованный список',
        'table': 'Таблица',
        'illustration': 'Иллюстрация',
        'table_reference': 'Ссылка на таблицу',
        'illustration_reference': 'Ссылка на иллюстрацию',
        'data_module_reference': 'Ссылка на модуль данных',
        'warning': 'Предупреждение'
    }

    with open(log_path, 'w', encoding='utf-8') as f:
        f.write("Анализ элементов документа\n")
        f.write(f"Сгенерировано: {datetime.datetime.now()}\n")
        f.write(f"Исходный файл: {document_path}\n")
        f.write("=" * 80 + "\n\n")

        for i, element in enumerate(elements, 1):
            element_type = element['type']
            type_name = element_type_names.get(element_type, element_type)

            f.write(f"Элемент {i}: {type_name}\n")

            # Position information
            if 'start_line' in element:
                f.write(f"  Позиция: строка {element['start_line']}, символ {element['start_char']}\n")
                if element.get('end_line') != element['start_line'] or element.get('end_char') != element['start_char']:
                    f.write(f"  Конец: строка {element.get('end_line', element['start_line'])}, символ {element.get('end_char', element['start_char'])}\n")

            # Content with more text shown
            content = element['content']
            if len(content) > 200:
                f.write(f"  Содержание: {content[:200]}...\n")
            else:
                f.write(f"  Содержание: {content}\n")

            if element.get('details'):
                f.write(f"  Детали: {element['details']}\n")

            f.write(f"  Пример XML: {element['xml_example']}\n")
            f.write("\n" + "-" * 80 + "\n\n")

    print(f"Лог элементов сохранен в: {log_path}")
    return log_path
