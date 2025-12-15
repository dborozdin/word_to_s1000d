"""
List parser for docx documents.
Extracts numbered and bulleted lists for S1000D format.
"""

import re
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches


def extract_lists(doc: Document) -> List[Dict[str, str]]:
    """
    Extract lists from document using OOXML formatting and heuristics.

    Args:
        doc: Docx document object

    Returns:
        List of dicts with list items and types
    """
    lists_data = []
    current_list = None
    list_items = []

    for para_idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        # Get list type using OOXML formatting
        list_type = _get_list_type_ooXML(paragraph)

        # If not detected by OOXML, try heuristic
        if not list_type:
            list_type = _get_list_type_heuristic(doc, para_idx)

        if list_type:
            if current_list is None or current_list['type'] != list_type:
                # Save previous list
                if current_list and list_items:
                    current_list['items'] = list_items.copy()
                    lists_data.append(current_list)

                # Start new list
                current_list = {'type': list_type, 'items': []}
                list_items = []

            # Add text to current list, cleaning markers if from heuristic
            clean_text = _clean_list_item_text(paragraph, list_type)
            if clean_text not in list_items:
                list_items.append(clean_text)
        else:
            # End current list if not empty
            if current_list and list_items:
                current_list['items'] = list_items.copy()
                lists_data.append(current_list)
                current_list = None
                list_items = []

    # Save final list
    if current_list and list_items:
        current_list['items'] = list_items
        lists_data.append(current_list)

    return lists_data


def _get_list_type_ooXML(paragraph) -> Optional[str]:
    """Get list type for paragraph based on OOXML formatting."""
    # Check if paragraph is formatted as a list in Word using OOXML numPr
    if hasattr(paragraph.paragraph_format, 'numPr') and paragraph.paragraph_format.numPr is not None:
        print(f"DEBUG: OOXML found numPr for paragraph: {repr(paragraph.text)}")
        # Get numId to determine list type
        num_id = paragraph.paragraph_format.numPr.numId
        print(f"DEBUG: num_id = {num_id}")
        if num_id is not None:
            # Access document numbering to determine if numbered or bulleted
            doc = paragraph.part.document
            numbering_part = doc.numbering_part
            if numbering_part is not None:
                try:
                    num = numbering_part.nums[num_id.val]
                    abstract_num = numbering_part.abstract_nums[num.abstract_num_id.val]
                    # Check lvl0 lvlText to determine list type - if contains % it's numbered
                    lvl0 = abstract_num.lvl0
                    if lvl0 is not None and hasattr(lvl0, 'lvlText') and lvl0.lvlText is not None:
                        lvl_text = lvl0.lvlText.val
                        if '%' in lvl_text:
                            return 'numbered_list'
                        else:
                            return 'unnumbered_list'
                    else:
                        # Fallback to numFmt if lvlText not available
                        if hasattr(abstract_num, 'num_style_link') and abstract_num.num_style_link is not None:
                            # If linked to a style, check if it's a numbered style
                            style_name = abstract_num.num_style_link.val.lower()
                            if 'number' in style_name or 'decimal' in style_name:
                                return 'numbered_list'
                            else:
                                return 'unnumbered_list'
                        else:
                            # Check lvl0 numFmt directly
                            lvl0 = abstract_num.lvl0
                            if lvl0 is not None and hasattr(lvl0, 'numFmt') and lvl0.numFmt is not None:
                                num_fmt = lvl0.numFmt.val.lower()
                                if num_fmt in ['decimal', 'lowerroman', 'upperroman', 'lowerletter', 'upperletter']:
                                    return 'numbered_list'
                                else:
                                    return 'unnumbered_list'
                            else:
                                return 'unnumbered_list'
                except (KeyError, AttributeError):
                    return 'unnumbered_list'
            else:
                return 'unnumbered_list'
        else:
            return 'unnumbered_list'

    return None


def _get_list_type_heuristic(doc: Document, para_idx: int) -> Optional[str]:
    """Get list type for paragraph using heuristic when OOXML formatting is not available."""
    unnumbered_markers = [
        '•', '◦', '▪', '▫', '⁃', '·', '-', '–', '—', '*', '†', '‡', '§'
    ]

    paragraph = doc.paragraphs[para_idx]

    text = paragraph.text#.strip()
    print(f"DEBUG: Heuristic checking para {para_idx}, text: {repr(text)}")
    if any(text.startswith(marker) for marker in unnumbered_markers):
        print(f"DEBUG: Found marker in text")
        return 'unnumbered_list'
    print(f"DEBUG: No marker found")

    # Additional heuristic: if text ends with ';' or '.', consider it a list item
    text_stripped = text.strip()
    if text_stripped.endswith((';', '.')) and not text_stripped.endswith(':'):
        return 'unnumbered_list'

    # Try simple approach: scan for introductory paragraph ending with ":" followed by list items
    result = _get_list_type_simple(doc, para_idx)
    if result:
        return result

    # Try colon-based approach as final fallback
    return _get_list_type_colon_intro(doc, para_idx)


def _clean_list_item_text(paragraph, list_type: str) -> str:
    """Clean list item text by removing markers and formatting."""
    import re
    text = paragraph.text.strip()

    if list_type == 'unnumbered_list':
        # Remove common unnumbered list markers
        unnumbered_markers = [
            '•', '◦', '▪', '▫', '⁃', '·', '-', '–', '—', '*', '†', '‡', '§'
        ]
        for marker in unnumbered_markers:
            if text.startswith(marker):
                text = text[len(marker):].strip()
                break
    elif list_type == 'numbered_list':
        # Remove numbered list markers using regex
        numbered_patterns = [
            r'^\d+\.\s*',  # 1., 2., 10.
            r'^\d+\)\s*',  # 1), 2), 10)
            r'^\(\d+\)\s*',  # (1), (2), (10)
            r'^\d+\s*[.)]\s*',  # 1 , 1), 2 , etc.
            r'^[a-zA-Z]\.\s*',  # a., b., A., B.
            r'^[a-zA-Z]\)\s*',  # a), b), A), B)
            r'^\([a-zA-Z]\)\s*',  # (a), (b), (A), (B)
            r'^[IVXLCDM]+\.\s*',  # I., II., III. (Roman numerals)
            r'^[ivxlcdm]+\.\s*',  # i., ii., iii. (Roman numerals)
            r'^[IVXLCDM]+\)\s*',  # I), II), III)
            r'^[ivxlcdm]+\)\s*',  # i), ii), iii)
        ]
        for pattern in numbered_patterns:
            text = re.sub(pattern, '', text, count=1)
            if text != paragraph.text.strip():  # If something was removed
                break

    return text.strip()


def _get_list_type_simple(doc: Document, para_idx: int) -> Optional[str]:
    """Simple approach: scan for introductory paragraph ending with ':' followed by list items until empty or ':'."""
    paragraphs = doc.paragraphs

    # Search backward for introductory paragraph ending with ":"
    intro_idx = None
    for idx in range(para_idx - 1, -1, -1):
        para = paragraphs[idx]
        text = para.text.strip()
        if not text:
            continue
        if text.endswith(':'):
            intro_idx = idx
            break

    if intro_idx is not None:
        # Check if current para is a list item
        current_idx = intro_idx + 1
        while current_idx < len(paragraphs):
            para = paragraphs[current_idx]
            text = para.text.strip()
            if not text:
                break
            if text.endswith(':'):
                break
            if current_idx == para_idx:
                return 'unnumbered_list'
            current_idx += 1

    return None


def _get_list_type_colon_intro(doc: Document, para_idx: int) -> Optional[str]:
    """
    Simple colon-based approach from tests/text_extract_lists.py.
    Detects lists by finding paragraphs ending with ':' followed by non-empty paragraphs.
    """
    paragraphs = doc.paragraphs
    
    # Search backward for introductory paragraph ending with ":"
    intro_idx = None
    for idx in range(para_idx - 1, -1, -1):
        para = paragraphs[idx]
        text = para.text.strip()
        if not text:
            # Empty line means we've passed the list boundary
            break
        if text.endswith(':'):
            intro_idx = idx
            break
    
    if intro_idx is None:
        return None
    
    # Check if current paragraph is within the list items after intro
    # List ends at empty line or at another line ending with ':'
    current_idx = intro_idx + 1
    list_items_indices = []
    
    while current_idx < len(paragraphs):
        para = paragraphs[current_idx]
        text = para.text.strip()
        
        # Empty line ends the list
        if not text:
            break
        
        # Another introductory line ends this list
        if text.endswith(':'):
            break
        
        list_items_indices.append(current_idx)
        current_idx += 1
    
    # If current paragraph is in the list items, return unnumbered_list
    if para_idx in list_items_indices and len(list_items_indices) >= 1:
        print(f"DEBUG: Colon-intro detected list item at para {para_idx}")
        return 'unnumbered_list'
    
    return None


def _get_list_type_third_approach(doc: Document, para_idx: int) -> Optional[str]:
    """Third approach: scan for introductory paragraph ending with ':' followed by list items."""
    paragraphs = doc.paragraphs

    # Find introductory paragraph ending with ':' before current paragraph
    intro_idx = None
    for idx in range(para_idx - 1, -1, -1):
        para = paragraphs[idx]
        text = para.text.strip()
        if not text:
            continue
        if text.endswith(':'):
            intro_idx = idx
            break

    if intro_idx is None:
        return None

    # Get the style of the introductory paragraph
    intro_para = paragraphs[intro_idx]
    intro_style = intro_para.style.name if intro_para.style else 'Normal'

    # Check if intro has header style
    if intro_style and ('heading' in intro_style.lower() or 'header' in intro_style.lower()):
        return None

    # Now check following non-empty paragraphs starting from intro_idx + 1
    list_items = []
    current_idx = intro_idx + 1

    while current_idx < len(paragraphs):
        para = paragraphs[current_idx]
        text = para.text.strip()

        # Stop at empty line
        if not text:
            break

        # Stop at header
        para_style = para.style.name if para.style else 'Normal'
        if para_style and ('heading' in para_style.lower() or 'header' in para_style.lower()):
            break

        # Stop at paragraph ending with ':'
        if text.endswith(':'):
            break

        # Check if same style as intro
        # if para_style != intro_style:
        #     break

        # This is a potential list item
        list_items.append(current_idx)
        current_idx += 1

    # If we found at least 2 items, and current para_idx is one of them, return unnumbered_list
    if len(list_items) >= 2 and para_idx in list_items:
        return 'unnumbered_list'

    return None


def convert_list_to_s1000d_randomlist(list_data: Dict[str, List[str]]) -> str:
    """
    Convert list data to S1000D randomList XML.

    Args:
        list_data: Dict with 'items' key

    Returns:
        XML string for randomList
    """
    if not list_data.get('items'):
        return ""

    items = []
    for item_text in list_data['items']:
        items.append(f"<listItem><para>{item_text}</para></listItem>")

    list_items_xml = ''.join(items)

    prefix = 'pf02' if list_data['type'] == 'unnumbered_list' else 'nfp01'

    return f'<randomList listItemPrefix="{prefix}">{list_items_xml}</randomList>'
