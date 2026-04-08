"""
Procedure module processor for tech cards (ТК / техкарты).
Parses Word documents and generates S1000D XML files conforming to proced.xsd.
"""

import os
import re
import configparser
from typing import Dict, List, Tuple
from docx import Document
from lxml import etree

from app_paths import long_path

# Import parser modules
from parsers.illustration_parser import extract_illustrations, ensure_missing_placeholders, copy_publication_logo
from parsers.elements_analyzer import analyze_document_elements, generate_elements_log, apply_reference_markup

# Import generators
from generators.s1000d_generator import S1000DGenerator, create_data_module_config
from generators.pm_generator import PMGenerator, create_pm_config, create_dm_ref_data

# Reuse helpers from descriptive processor
from processing_scripts.descriptive_processor import (
    normalize_title_case,
    extract_document_title,
    extract_organization_from_document,
)


# ---------------------------------------------------------------------------
# Preliminary requirements parsing
# ---------------------------------------------------------------------------

# Keywords that indicate equipment tables
EQUIPMENT_KEYWORDS = [
    'оснастка', 'инструмент', 'оборудование', 'средства обеспечения',
    'средство обеспечения', 'приспособление', 'наименование средств',
]

# Keywords that indicate supply/material tables
SUPPLY_KEYWORDS = [
    'расходные материалы', 'расходный материал', 'материалы', 'горюче-смазочные',
    'гсм', 'наименование материал',
]

# Keywords that indicate the start of procedural steps
STEP_START_KEYWORDS = [
    'порядок выполнения', 'последовательность операций', 'содержание операции',
    'технологическая последовательность', 'порядок работ',
]

# Keywords for section headers within procedures (nested steps)
SECTION_HEADER_KEYWORDS = [
    'демонтаж', 'монтаж', 'разборка', 'сборка', 'подготовка',
    'проверка', 'установка', 'снятие', 'замена', 'регулировка',
    'настройка', 'испытание', 'осмотр', 'обслуживание',
]

# Keywords for safety warnings in preliminary section
WARNING_KEYWORDS = ['внимание', 'осторожно', 'предупреждение', 'опасно']
NOTE_KEYWORDS = ['примечание', 'примечания']

# XML namespaces for VML textbox table extraction
VML_NS = 'urn:schemas-microsoft-com:vml'
WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


# ---------------------------------------------------------------------------
# VML textbox table extraction
# ---------------------------------------------------------------------------

def _extract_tables_from_vml_textboxes(doc: Document) -> List[Dict]:
    """
    Extract tables embedded inside VML floating textboxes.

    In .doc files converted to .docx, tables for equipment/supplies are often
    inside VML shapes: <w:p>/<w:r>/<w:pict>/<v:shape>/<v:textbox>/<w:txbxContent>/<w:tbl>
    These are invisible to doc.tables which only returns top-level body tables.

    Returns:
        List of table dicts: [{'rows': [['cell1', 'cell2', ...], ...]}]
    """
    tables = []

    # Search entire document body XML for VML shapes containing tables
    for element in doc.element.body:
        # Look for v:shape elements anywhere in this body child
        for shape in element.iter('{%s}shape' % VML_NS):
            for textbox in shape.iter('{%s}textbox' % VML_NS):
                for txbx_content in textbox.iter('{%s}txbxContent' % WORD_NS):
                    for tbl in txbx_content.iter('{%s}tbl' % WORD_NS):
                        table_data = _parse_vml_table_element(tbl)
                        if table_data and table_data['rows']:
                            tables.append(table_data)

    return tables


def _parse_vml_table_element(tbl_element) -> Dict:
    """
    Parse a <w:tbl> lxml element into a dict of rows/cells.

    Each cell is represented as a list of paragraph texts (preserving
    paragraph boundaries, since equipment/supply items are individual paragraphs
    within a single cell).

    Args:
        tbl_element: lxml element for <w:tbl>

    Returns:
        {'rows': [[['para1', 'para2'], ['para1'], ...], ...]}
        Each cell is a list of paragraph strings.
    """
    rows = []

    for tr in tbl_element.findall('{%s}tr' % WORD_NS):
        row_cells = []
        for tc in tr.findall('{%s}tc' % WORD_NS):
            cell_paras = []
            for p in tc.findall('{%s}p' % WORD_NS):
                para_text = ''
                for t in p.iter('{%s}t' % WORD_NS):
                    if t.text:
                        para_text += t.text
                para_text = para_text.strip()
                if para_text:
                    cell_paras.append(para_text)
            row_cells.append(cell_paras)
        if row_cells:
            rows.append(row_cells)

    return {'rows': rows}


def parse_preliminary_requirements(doc: Document, elements: List[Dict]) -> Dict:
    """
    Parse preliminary requirements from the beginning of the tech card document.

    Extracts equipment, supplies, safety notes/warnings from tables and paragraphs
    before the procedural steps begin.

    Returns:
        Dict with keys: support_equips, supplies, spares, safety_notes,
        safety_warnings, safety_cautions, prelim_end_index
    """
    result = {
        'support_equips': [],
        'supplies': [],
        'spares': [],
        'safety_notes': [],
        'safety_warnings': [],
        'safety_cautions': [],
        'prelim_end_index': 0,  # element index where prelim ends
    }

    # Extract tables from VML textboxes (common in .doc → .docx converted files)
    vml_tables = _extract_tables_from_vml_textboxes(doc)
    if vml_tables:
        print(f"  Found {len(vml_tables)} table(s) in VML textboxes")
        tables_classified = _classify_tables_from_dicts(vml_tables)
    else:
        # Fallback: try top-level doc.tables (standard .docx files)
        top_level_tables = [_docx_table_to_dict(t) for t in doc.tables]
        tables_classified = _classify_tables_from_dicts(top_level_tables)
    print(f"  Classified: equipment={len(tables_classified['equipment'])}, supplies={len(tables_classified['supplies'])}")
    result['support_equips'] = tables_classified.get('equipment', [])
    result['supplies'] = tables_classified.get('supplies', [])
    result['spares'] = tables_classified.get('spares', [])

    # Find where procedural steps start — scan elements for the boundary
    step_start_idx = _find_step_start_index(elements)
    result['prelim_end_index'] = step_start_idx

    # Extract safety notes and warnings from paragraphs before step_start_idx
    for elem in elements[:step_start_idx]:
        elem_type = elem.get('type', '')
        content = elem.get('content', '').strip()
        if not content:
            continue

        content_lower = content.lower()

        if elem_type == 'warning' or any(kw in content_lower for kw in WARNING_KEYWORDS):
            # Distinguish between warning and caution
            if 'осторожно' in content_lower:
                result['safety_cautions'].append(_clean_safety_text(content))
            else:
                result['safety_warnings'].append(_clean_safety_text(content))
        elif any(kw in content_lower for kw in NOTE_KEYWORDS):
            # It's a note marker line — the actual text may follow
            note_text = _extract_note_text(content)
            if note_text:
                result['safety_notes'].append(note_text)
        elif elem_type == 'paragraph' and _looks_like_note(content, elements, elem):
            result['safety_notes'].append(content)

    return result


def _docx_table_to_dict(table) -> Dict:
    """Convert a python-docx Table object to the same format as _parse_vml_table_element."""
    rows = []
    for row in table.rows:
        row_cells = []
        for cell in row.cells:
            # Split cell text by lines (approximate paragraph boundaries)
            lines = [line.strip() for line in cell.text.split('\n') if line.strip()]
            row_cells.append(lines)
        rows.append(row_cells)
    return {'rows': rows}


def _classify_tables_from_dicts(tables: List[Dict]) -> Dict[str, List]:
    """
    Classify VML tables into equipment, supplies, spares.

    Tech card tables use a column-based layout where the header row identifies
    column types (e.g. "Инструмент и приспособления" | "Расходные материалы"),
    and data cells contain multiple items as separate paragraphs within one cell.

    Args:
        tables: List of table dicts from _parse_vml_table_element().
                Each cell is a list of paragraph strings.

    Returns:
        Dict with 'equipment', 'supplies', 'spares' lists of item dicts.
    """
    classified = {'equipment': [], 'supplies': [], 'spares': []}

    for table_dict in tables:
        rows = table_dict.get('rows', [])
        if len(rows) < 2:
            continue

        # Identify column types from header row(s)
        header_row = rows[0]
        col_types = {}  # column index -> type string

        for col_idx, cell_paras in enumerate(header_row):
            cell_text = ' '.join(cell_paras).lower()
            if any(kw in cell_text for kw in EQUIPMENT_KEYWORDS):
                col_types[col_idx] = 'equipment'
            elif any(kw in cell_text for kw in SUPPLY_KEYWORDS):
                col_types[col_idx] = 'supplies'
            elif 'запасные части' in cell_text or 'запчаст' in cell_text:
                col_types[col_idx] = 'spares'

        if not col_types:
            continue

        # Extract items from data rows (skip header)
        for row in rows[1:]:
            for col_idx, item_type in col_types.items():
                if col_idx >= len(row):
                    continue
                cell_paras = row[col_idx]
                for para_text in cell_paras:
                    name = para_text.strip()
                    if name and len(name) > 3:
                        classified[item_type].append({'name': name})

    return classified


def _find_step_start_index(elements: List[Dict]) -> int:
    """
    Find the element index where procedural steps begin.

    Heuristic:
    1. Look for explicit step-start headers
    2. Look for numbered paragraphs after tables
    3. Default: after the last table element
    """
    last_table_idx = 0

    for idx, elem in enumerate(elements):
        elem_type = elem.get('type', '')
        content = elem.get('content', '').strip().lower()

        # Track last table position
        if elem_type == 'table':
            last_table_idx = idx + 1

        # Check for explicit step-start keywords
        if any(kw in content for kw in STEP_START_KEYWORDS):
            return idx + 1  # steps start after this header

    # If no explicit header found, look for first numbered element,
    # section header, or imperative paragraph after the last table
    for idx in range(last_table_idx, len(elements)):
        elem = elements[idx]
        elem_type = elem.get('type', '')
        content = elem.get('content', '').strip()

        if elem_type in ('numbered_paragraph_header', 'numbered_list'):
            return idx

        # Check for section headers (Демонтаж, Монтаж, etc.)
        if content and _is_section_header(content, elem_type):
            return idx

        # Check for imperative verb (common step start pattern)
        if elem_type == 'paragraph' and content and _is_imperative(content):
            return idx

    # Fallback: everything after tables
    return last_table_idx


def _is_imperative(text: str) -> bool:
    """Check if text starts with a Russian imperative verb (common in procedures)."""
    imperative_prefixes = [
        'откройте', 'закройте', 'снимите', 'установите', 'проверьте',
        'убедитесь', 'отсоедините', 'подсоедините', 'расконтрите',
        'законтрите', 'ослабьте', 'заверните', 'произведите', 'подключите',
        'отключите', 'включите', 'выключите', 'слейте', 'заправьте',
        'зарядите', 'разрядите', 'стравите', 'прокачайте', 'закачайте',
        'удалите', 'нанесите', 'протрите', 'расстыкуйте', 'состыкуйте',
        'расконсервируйте', 'выполните', 'выдержите', 'дозарядите',
        'доразрядите', 'уменьшите', 'создайте',
    ]
    first_word = text.split()[0].lower().rstrip('.,;:') if text.split() else ''
    return first_word in imperative_prefixes


def _clean_safety_text(text: str) -> str:
    """Remove safety prefix markers from text."""
    prefixes = ['ВНИМАНИЕ!', 'ВНИМАНИЕ:', 'ВНИМАНИЕ', 'ОСТОРОЖНО!', 'ОСТОРОЖНО:', 'ОСТОРОЖНО',
                'ПРЕДУПРЕЖДЕНИЕ!', 'ПРЕДУПРЕЖДЕНИЕ:', 'ПРЕДУПРЕЖДЕНИЕ']
    result = text.strip()
    for prefix in prefixes:
        if result.upper().startswith(prefix):
            result = result[len(prefix):].strip()
            break
    return result


def _extract_note_text(text: str) -> str:
    """Extract note text after 'Примечание' prefix."""
    patterns = [
        r'[Пп]римечани[ея]\s*[.:–-]\s*(.+)',
        r'[Пп]римечани[ея]\s+(.+)',
    ]
    for pattern in patterns:
        match = re.match(pattern, text, re.DOTALL)
        if match:
            return match.group(1).strip()
    return ''


def _looks_like_note(content: str, elements: List[Dict], current_elem: Dict) -> bool:
    """Check if a paragraph looks like a safety note based on context."""
    # Notes are often italicized or have specific patterns
    return False


# ---------------------------------------------------------------------------
# List grouping for procedural steps
# ---------------------------------------------------------------------------

import re as _re

_LIST_ITEM_RE = _re.compile(
    r'^[\-\u2013\u2014\u2212\u2022\u25CF\u25CB\u2023\u25AA\u25AB\u2043]\s+'
)


def _is_list_item_by_content(step: Dict) -> bool:
    """Detect list item by type OR by text starting with dash/bullet."""
    step_type = step.get('type', '')
    if step_type in ('unnumbered_list', 'nested_unnumbered_list'):
        return True
    # Fallback: check text content (handles cases where apply_reference_markup
    # changed type from unnumbered_list to paragraph)
    text = step.get('text', '').strip()
    if _LIST_ITEM_RE.match(text):
        return True
    # Semicolon-ending short text after a colon-intro is also a list item
    if text.endswith(';') and len(text) < 120:
        return True
    return False


def _group_list_items(steps: List[Dict]) -> List[Dict]:
    """Group consecutive list items into the preceding step.

    Before: [step_intro "Выберите:", list "от промышленной", list "от автономного"]
    After:  [step_intro "Выберите:" with list_items=["от промышленной", "от автономного"]]
    """
    if not steps:
        return steps

    result = []
    i = 0
    while i < len(steps):
        step = steps[i]

        # Check if NEXT elements are list items
        if i + 1 < len(steps) and _is_list_item_by_content(steps[i + 1]):
            # Collect consecutive list items
            list_items = []
            j = i + 1
            while j < len(steps) and _is_list_item_by_content(steps[j]):
                list_items.append({'text': steps[j]['text']})
                j += 1

            # Attach list items to current step
            step = dict(step)  # copy
            step['list_items'] = list_items
            result.append(step)
            i = j
            continue

        # Current step IS a list item but no preceding intro
        if _is_list_item_by_content(step):
            # Check if previous result step can absorb it
            if result and 'list_items' not in result[-1]:
                prev = result[-1]
                prev['list_items'] = [{'text': step['text']}]
                j = i + 1
                while j < len(steps) and _is_list_item_by_content(steps[j]):
                    prev['list_items'].append({'text': steps[j]['text']})
                    j += 1
                i = j
                continue

        result.append(step)
        i += 1

    return result


# ---------------------------------------------------------------------------
# Procedural steps parsing
# ---------------------------------------------------------------------------

def parse_procedural_steps(elements: List[Dict], start_index: int) -> List[Dict]:
    """
    Parse procedural steps from elements starting at start_index.

    Detects section headers (e.g. "Демонтаж", "Монтаж") and creates
    nested step structures.

    Returns:
        List of step dicts: {'text': str, 'substeps': List[Dict]}
    """
    step_elements = elements[start_index:]
    if not step_elements:
        return []

    # First pass: collect all step-like elements
    raw_steps = []
    for elem in step_elements:
        elem_type = elem.get('type', '')
        content = elem.get('content', '').strip()
        if not content:
            continue

        # Skip table and illustration elements in procedure section
        if elem_type in ('table', 'table_reference', 'illustration',
                         'illustration_reference'):
            continue

        # Determine if this is a section header
        is_section = _is_section_header(content, elem_type)

        raw_steps.append({
            'text': content,
            'type': elem_type,
            'is_section': is_section,
            'substeps': [],
        })

    # Group consecutive list items into parent step's list_items
    raw_steps = _group_list_items(raw_steps)

    # Second pass: nest steps under section headers
    if not raw_steps:
        return []

    # Check if any section headers exist
    has_sections = any(s['is_section'] for s in raw_steps)

    if not has_sections:
        # Flat structure: all steps at top level
        return [{'text': s['text'], 'substeps': [],
                 'list_items': s.get('list_items', [])} for s in raw_steps]

    # Nested structure: group steps under their section headers
    result = []
    current_section = None

    for step in raw_steps:
        if step['is_section']:
            # Start a new section
            if current_section is not None:
                result.append(current_section)
            current_section = {
                'text': step['text'],
                'substeps': [],
                'list_items': step.get('list_items', []),
            }
        else:
            if current_section is not None:
                current_section['substeps'].append({
                    'text': step['text'],
                    'substeps': [],
                    'list_items': step.get('list_items', []),
                })
            else:
                # Steps before any section header — add at top level
                result.append({'text': step['text'], 'substeps': [],
                               'list_items': step.get('list_items', [])})

    # Don't forget the last section
    if current_section is not None:
        result.append(current_section)

    return result


def _is_section_header(text: str, elem_type: str) -> bool:
    """
    Determine if a text is a procedure section header (like Демонтаж, Монтаж).

    Section headers are:
    - Short (< 80 chars)
    - Don't end with period
    - Contain known section keywords or are typed as header/numbered_paragraph_header
    """
    text_stripped = text.strip()

    # Too long to be a section header
    if len(text_stripped) > 80:
        return False

    # Check element type
    if elem_type in ('header', 'numbered_paragraph_header'):
        return True

    # Check for known section keywords
    text_lower = text_stripped.lower()
    # Must be relatively short and match a keyword
    if len(text_stripped) < 60 and not text_stripped.endswith('.'):
        for keyword in SECTION_HEADER_KEYWORDS:
            if keyword in text_lower:
                return True

    return False


# ---------------------------------------------------------------------------
# Main processing function
# ---------------------------------------------------------------------------

def process_procedure_document(doc_path: str, output_dir: str, llm_config: Dict = None,
                                dm_code_override: Dict = None, tech_name_override: str = None,
                                info_name_override: str = None, skip_pmc: bool = False,
                                graphic_ident_prefix: str = None):
    """
    Process procedure document (техкарта): parse, map to S1000D procedure structure,
    generate XML conforming to proced.xsd.

    Args:
        doc_path: Path to docx document
        output_dir: Output directory for generated files
        llm_config: Optional LLM configuration
        dm_code_override: Optional DMC code dict override
        tech_name_override: Optional techName override (from folder name)
        info_name_override: Optional infoName override (from folder name)
        skip_pmc: If True, skip PMC generation (for batch mode)
        graphic_ident_prefix: Optional prefix for graphic entity naming

    Returns:
        Tuple of (dm_refs list, illustrations dict)
    """
    print(f"Processing procedure document: {doc_path}")

    # Load document
    doc = Document(long_path(doc_path))

    # Extract organization from headers/footers (reused)
    organization = extract_organization_from_document(doc)
    print(f"Extracted organization: {organization}")

    # Extract document title (reused)
    document_title = extract_document_title(doc)
    print(f"Extracted document title: {document_title}")

    # Extract illustrations
    print("Extracting illustrations...")
    illustrations, illustration_positions = extract_illustrations(
        doc, output_dir, graphic_ident_prefix=graphic_ident_prefix
    )
    copy_publication_logo(output_dir)

    # Analyze document elements
    print("Analyzing document elements...")
    elements = analyze_document_elements(
        doc, illustrations, illustration_positions,
        llm_config=llm_config, graphic_ident_prefix=graphic_ident_prefix
    )

    # Apply user reference markup (or fall back to overrides) if available
    if dm_code_override:
        from parsers.dmc_parser import dm_code_to_string
        dmc_str = dm_code_to_string(dm_code_override)
        elements = apply_reference_markup(elements, dmc_str)

    # Generate elements log
    elements_log_path = generate_elements_log(doc_path, elements, output_dir)

    # --- Procedure-specific parsing ---

    # 1. Parse preliminary requirements (equipment, supplies, safety)
    print("Parsing preliminary requirements...")
    prelim_data = parse_preliminary_requirements(doc, elements)
    print(f"  Equipment items: {len(prelim_data['support_equips'])}")
    print(f"  Supply items: {len(prelim_data['supplies'])}")
    print(f"  Safety notes: {len(prelim_data['safety_notes'])}")
    print(f"  Safety warnings: {len(prelim_data['safety_warnings'])}")

    # 2. Parse procedural steps
    print("Parsing procedural steps...")
    steps = parse_procedural_steps(elements, prelim_data['prelim_end_index'])
    total_steps = sum(1 + len(s.get('substeps', [])) for s in steps)
    print(f"  Procedural steps: {total_steps} (top-level: {len(steps)})")

    # 3. Build procedure content structure for generator
    procedure_content = {
        'preliminary_rqmts': prelim_data,
        'procedural_steps': steps,
    }

    # Determine DM code
    dm_code = dm_code_override if dm_code_override else _get_default_procedure_dm_code()

    # Determine titles
    effective_tech_name = tech_name_override or document_title
    effective_info_name = info_name_override or ''

    # Create DM config
    dm_config = create_data_module_config(
        effective_tech_name,
        effective_info_name,
        dm_code,
        procedure_content,
        enterprise_name=organization,
        originator_name=organization,
    )

    # Generate procedure XML
    generator = S1000DGenerator()
    figure_info = []

    filepath = generator.generate_procedure_module(
        dm_config, output_dir, illustrations, figure_info
    )

    # Generate placeholder images for missing illustrations
    ensure_missing_placeholders(figure_info, output_dir)

    # Collect DM ref for PM generation
    dm_ref = create_dm_ref_data(dm_code, effective_tech_name, effective_info_name)
    dm_refs = [dm_ref]

    generated_files = [filepath]

    print(f"\nGenerated {len(generated_files)} files:")
    for fp in generated_files:
        print(f"  - {os.path.basename(fp)}")
    print(f"Elements log: {os.path.basename(elements_log_path)}")

    return dm_refs, illustrations


def _get_default_procedure_dm_code() -> Dict:
    """Return default DM code for a procedure module."""
    return {
        'modelIdentCode': 'S5',
        'systemDiffCode': 'A',
        'systemCode': '029',
        'subSystemCode': '1',
        'subSubSystemCode': '1',
        'assyCode': '01',
        'disassyCode': '00',
        'disassyCodeVariant': 'A',
        'infoCode': '242',
        'infoCodeVariant': 'A',
        'itemLocationCode': 'A',
    }
